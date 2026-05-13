"""
CAM Lease Analyzer Web App — FastAPI Application

Domain-agnostic routes + lease-specific routes, static file serving, CORS.
"""

import hashlib
import json
import logging
import os
import shutil
import threading
import time
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from starlette.staticfiles import NotModifiedResponse
from starlette.responses import Response

# config.py sets up sys.path for cam imports
from app.config import get_config, email_configured, LEASE_DIR, APP_VERSION, GIT_SHA
from app import job_manager

# ── Logging ──
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(name)s] %(levelname)s: %(message)s",
)
logger = logging.getLogger(__name__)

# ── CAM Knowledge Base — loaded once at startup, injected into all chat prompts ──
_KNOWLEDGE_PATH = Path(__file__).parent / "cam_knowledge.txt"
CAM_KNOWLEDGE = _KNOWLEDGE_PATH.read_text(encoding="utf-8") if _KNOWLEDGE_PATH.exists() else ""


def _normalize_followups(items: Any, limit: int = 3) -> List[str]:
    """Return up to `limit` clean follow-up question strings."""
    if not isinstance(items, list):
        return []
    cleaned: List[str] = []
    for item in items:
        if not isinstance(item, str):
            continue
        text = " ".join(item.strip().split())
        if not text:
            continue
        if len(text) > 120:
            text = text[:117].rstrip() + "..."
        cleaned.append(text)
        if len(cleaned) >= limit:
            break
    return cleaned


def _parse_chat_json_payload(raw_text: str) -> Optional[Dict[str, Any]]:
    """Parse a JSON object from a model response, tolerating fenced blocks."""
    if not raw_text:
        return None
    text = raw_text.strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:].strip()
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return None
    try:
        parsed = json.loads(text[start:end + 1])
        return parsed if isinstance(parsed, dict) else None
    except json.JSONDecodeError:
        return None


# Step 262: perspective addendum for chat SCOPE block. Tells the model whose
# interests to frame coverage gaps and risks against. The string is appended
# to both Mode A and Mode C scope blocks.
def _build_perspective_addendum(perspective: str) -> str:
    """Return a perspective-aware paragraph for the chat SCOPE block.

    Note: the upstream coverage assessor uses tenant-leaning regex patterns to
    classify states like 'covered_unfavorable'. The addendum below tells the
    chat model how to RE-FRAME those classifications when the user is reviewing
    on behalf of the landlord or as a neutral party. Coverage state values
    themselves are unchanged.
    """
    p = (perspective or "tenant").lower()
    if p == "landlord":
        return (
            "PERSPECTIVE: This analysis is being reviewed from the LANDLORD's perspective.\n"
            "Frame all coverage gaps, exposure statements, and recommendations against\n"
            "landlord interests. The underlying coverage state classifications (e.g.\n"
            "'covered_unfavorable') were generated using rules that lean tenant-protective\n"
            "— a clause flagged 'covered_unfavorable' may actually be FAVORABLE to the\n"
            "landlord (e.g. waived audit rights, sole-discretion consent, narrow force\n"
            "majeure). When that's the case, describe the landlord's UPSIDE plainly. When\n"
            "a 'missing' or 'partial' element instead exposes the landlord (e.g. no\n"
            "late-fee mechanism, no acceleration on default, no removal obligation at\n"
            "expiration), describe that landlord-side risk directly. Do not advocate for\n"
            "tenant interests in your answer. When asked 'what should I push for',\n"
            "recommend changes that protect the landlord.\n\n"
        )
    if p == "neutral":
        return (
            "PERSPECTIVE: This analysis is being reviewed NEUTRALLY (commercially reasonable).\n"
            "Do not advocate for either party. For each gap or unfavorable clause, identify\n"
            "which party benefits and which is exposed. If a clause is one-sided in either\n"
            "direction, name the imbalance plainly. If it's mutual, say so. The underlying\n"
            "coverage state classifications were generated using rules that lean\n"
            "tenant-protective — read past that bias and re-frame each finding even-handedly.\n"
            "When asked 'what should I push for', describe a balanced compromise rather than\n"
            "a one-sided ask.\n\n"
        )
    # Default: tenant. Matches the upstream rule set's bias, so no re-framing needed.
    return (
        "PERSPECTIVE: This analysis is being reviewed from the TENANT's perspective.\n"
        "Frame coverage gaps, exposure statements, and recommendations against tenant\n"
        "interests. When asked 'what should I push for', recommend changes that protect\n"
        "the tenant.\n\n"
    )


def _call_chat_with_followups(call_llm, provider: str, system_prompt: str, user_prompt: str, temperature: float = 0.3) -> Dict[str, Any]:
    """Call the model and ask for a main response plus suggested follow-up questions."""
    format_prompt = (
        f"{user_prompt}\n\n"
        "Return valid JSON only with this shape:\n"
        "{\n"
        '  "response": "main answer as a plain string",\n'
        '  "suggested_followups": ["short follow-up question", "short follow-up question", "short follow-up question"]\n'
        "}\n"
        "Rules for suggested_followups:\n"
        "- Provide 2 or 3 short, natural follow-up questions.\n"
        "- Make them useful for the current screen and answer.\n"
        "- Do not invent UI elements or unsupported actions.\n"
        "- Keep each follow-up under 12 words when possible.\n"
        "- If no good follow-ups exist, return an empty array.\n"
    )
    result = call_llm(
        provider=provider,
        system_prompt=system_prompt,
        user_prompt=format_prompt,
        temperature=temperature,
    )
    raw_content = result.get("content", "")
    payload = _parse_chat_json_payload(raw_content)
    if payload:
        response_text = str(payload.get("response", "") or "").strip()
        followups = _normalize_followups(payload.get("suggested_followups", []))
        if response_text:
            return {
                "content": response_text,
                "suggested_followups": followups,
                "model": result.get("model", ""),
                "raw_content": raw_content,
            }
    return {
        "content": raw_content,
        "suggested_followups": [],
        "model": result.get("model", ""),
        "raw_content": raw_content,
    }

# ── FastAPI app ──
app = FastAPI(
    title="CAM Lease Analyzer",
    description="Commercial lease deviation analysis powered by CAM",
    version="0.1.0",
)

# ── CORS (open for MVP) ──
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Startup ──

@app.on_event("startup")
def startup():
    from datetime import datetime
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"\n{'='*60}", flush=True)
    print(f"  CAM LEASE ANALYZER — SERVER STARTED", flush=True)
    print(f"  {now}", flush=True)
    print(f"  version={APP_VERSION}  sha={GIT_SHA}", flush=True)
    print(f"{'='*60}\n", flush=True)
    logger.info(f"app_startup version={APP_VERSION} sha={GIT_SHA}")

    config = get_config()

    # Load API keys for pipeline
    from cam.core.config import find_and_load_env
    find_and_load_env()

    # Ensure directories exist
    Path(config["UPLOAD_DIR"]).mkdir(parents=True, exist_ok=True)
    Path(config["RESULTS_DIR"]).mkdir(parents=True, exist_ok=True)

    # Load any persisted completed jobs
    loaded = job_manager.load_persisted_jobs()
    logger.info(f"Startup: {loaded} persisted job(s) loaded")

    # Clean up any expired jobs from previous runs
    cleaned = job_manager.cleanup_expired_jobs()
    if cleaned:
        logger.info(f"Startup cleanup: removed {cleaned} expired job(s)")

    # Start periodic cleanup thread (every 30 minutes)
    def _periodic_cleanup():
        while True:
            time.sleep(1800)
            try:
                job_manager.cleanup_expired_jobs()
            except Exception as e:
                logger.error(f"Periodic cleanup error: {e}")

    cleanup_thread = threading.Thread(target=_periodic_cleanup, daemon=True, name="job-cleanup")
    cleanup_thread.start()
    expiry_min = config.get('JOB_EXPIRY_MINUTES', 1440)
    expiry_display = f"{expiry_min // 60}h" if expiry_min >= 60 else f"{expiry_min}min"
    logger.info(f"Job expiry: {expiry_display} from completion, cleanup every 30 min")

    # Log email status
    if email_configured(config):
        logger.info("Email notifications: CONFIGURED")
    else:
        logger.info(
            "Email notifications: NOT CONFIGURED. "
            "Set SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASSWORD env vars. See .env.example."
        )

    # Validate access code
    if not config["ACCESS_CODE"]:
        logger.warning("ACCESS_CODE not set — access gate disabled")
    else:
        logger.info(f"Access code: configured ({len(config['ACCESS_CODE'])} chars)")


# ── Static files — no-cache for JS/CSS so version bumps take effect immediately ──
static_dir = LEASE_DIR / "static"
if static_dir.exists():
    class NoCacheStaticFiles(StaticFiles):
        async def get_response(self, path, scope):
            response = await super().get_response(path, scope)
            if path.endswith(('.js', '.css')):
                response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate'
                response.headers['Pragma'] = 'no-cache'
                response.headers['Expires'] = '0'
            return response
    app.mount("/static", NoCacheStaticFiles(directory=str(static_dir)), name="static")


# ══════════════════════════════════════════════════════════════════════════════
# GENERIC ROUTES (domain-agnostic — work for any CAM adapter)
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/")
def serve_index():
    """Serve the frontend index.html."""
    index_path = static_dir / "index.html"
    if index_path.exists():
        return FileResponse(
            str(index_path),
            headers={"Cache-Control": "no-cache, no-store, must-revalidate"},
        )
    return JSONResponse({"message": "CAM Lease Analyzer — frontend not yet built"})


@app.post("/api/auth/verify")
def verify_access_code(body: dict):
    """Verify the shared access code."""
    config = get_config()
    submitted = body.get("access_code", "")

    if not config["ACCESS_CODE"]:
        # No access code configured — allow all
        return {"valid": True}

    if submitted == config["ACCESS_CODE"]:
        return {"valid": True}

    raise HTTPException(status_code=401, detail="Invalid access code")


@app.get("/api/jobs/{job_id}")
def get_job_status(job_id: str):
    """Return job status dict. Returns 410 Gone if expired.
    Uses get_job_snapshot() for thread-safe serialization — the live
    job dict is mutated by background threads during processing."""
    job = job_manager.get_job_snapshot(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    # Check expiration
    if job.get("status") == "completed" and job_manager.is_job_expired(job):
        job_manager.cleanup_expired_jobs()
        return JSONResponse(
            status_code=410,
            content={"detail": "This analysis has expired. Results are no longer available."},
        )

    # Prevent browser from caching poll responses
    return JSONResponse(
        content=json.loads(json.dumps(job, default=str)),
        headers={"Cache-Control": "no-store"},
    )


@app.post("/api/jobs/{job_id}/cancel")
def cancel_job(job_id: str):
    """Request cancellation of a running job.

    Only valid when job status is 'processing'. The pipeline will stop
    between stages/tenants. Any completed tenant results are preserved.
    """
    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job.get("status") != "processing":
        raise HTTPException(
            status_code=400,
            detail=f"Cannot cancel job with status '{job.get('status')}'. Only 'processing' jobs can be cancelled.",
        )

    job_manager.request_cancel(job_id)
    return {"status": "cancel_requested"}


@app.patch("/api/jobs/{job_id}/email")
async def update_job_email(job_id: str, body: dict):
    """Update the notification email for a running or completed job."""
    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    email = body.get("email", "").strip()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Valid email address required")

    with job_manager._jobs_lock:
        if job_id in job_manager._jobs:
            job_manager._jobs[job_id]["email"] = email

    job_manager.save_job_results(job_id)
    return {"ok": True, "email": email}


@app.post("/api/send-results-link")
async def send_results_link(body: dict):
    """Send the results URL to an email address (mobile convenience)."""
    from app.notifications import _send_email

    email = (body.get("email") or "").strip()
    url = (body.get("url") or "").strip()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Valid email address required")
    if not url:
        raise HTTPException(status_code=400, detail="URL required")

    subject = "Your CAM Lease Analysis Results"
    text = (
        f"Here is the link to your lease analysis results:\n\n"
        f"{url}\n\n"
        f"Open this link on a desktop browser for the best experience.\n"
    )
    ok = _send_email(email, subject, text)
    if not ok:
        raise HTTPException(status_code=500, detail="Failed to send email")
    return {"ok": True}


@app.get("/api/jobs/{job_id}/results")
def get_job_results(job_id: str):
    """Return available results JSON for tenants. 404 only if no tenant results exist yet."""
    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    # Check expiration
    if job.get("status") == "completed" and job_manager.is_job_expired(job):
        job_manager.cleanup_expired_jobs()
        return JSONResponse(
            status_code=410,
            content={"detail": "This analysis has expired. Results are no longer available."},
        )

    tenants = job.get("input_config", {}).get("tenants", [])
    any_results_available = False
    results = []
    for i, tenant in enumerate(tenants):
        result_path = tenant.get("result_path")
        if result_path and Path(result_path).exists():
            try:
                data = json.loads(Path(result_path).read_text(encoding="utf-8"))
                annotated = tenant.get("annotated_path")
                comparison = tenant.get("comparison_view_path")
                any_results_available = True
                results.append({
                    "tenant_index": i,
                    "filename": tenant["filename"],
                    "status": tenant["status"],
                    "results": data,
                    "has_annotated": bool(annotated and Path(annotated).exists()),
                    "has_comparison_view": bool(comparison and Path(comparison).exists()),
                })
            except (json.JSONDecodeError, OSError) as e:
                results.append({
                    "tenant_index": i,
                    "filename": tenant["filename"],
                    "status": "error",
                    "error": str(e),
                })
        else:
            results.append({
                "tenant_index": i,
                "filename": tenant["filename"],
                "status": tenant.get("status", "unknown"),
                "error": tenant.get("error"),
            })

    if not any_results_available and job["status"] not in ("completed", "cancelled"):
        raise HTTPException(status_code=404, detail="Job not yet complete")

    return {"job_id": job_id, "tenants": results}


@app.get("/api/jobs")
def list_jobs(email: Optional[str] = None):
    """List all jobs, optionally filtered by email."""
    return job_manager.list_jobs(email=email)


@app.post("/api/jobs/{job_id}/feedback")
def submit_feedback(job_id: str, body: dict):
    """Append feedback for a specific finding."""
    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    tenant_index = body.get("tenant_index")
    provision_id = body.get("provision_id")
    assessment = body.get("assessment")

    if tenant_index is None or not provision_id or not assessment:
        raise HTTPException(
            status_code=400,
            detail="Required fields: tenant_index, provision_id, assessment",
        )

    if assessment not in ("agree", "disagree", "unsure"):
        raise HTTPException(
            status_code=400,
            detail="assessment must be 'agree', 'disagree', or 'unsure'",
        )

    saved = job_manager.add_feedback(
        job_id=job_id,
        tenant_index=tenant_index,
        provision_id=provision_id,
        assessment=assessment,
        notes=body.get("notes"),
    )

    if not saved:
        raise HTTPException(status_code=500, detail="Failed to save feedback")

    return {"saved": True}


# ══════════════════════════════════════════════════════════════════════════════
# RESOLUTION WORKFLOW (079)
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/jobs/{job_id}/resolutions")
def get_resolutions(job_id: str):
    """Return all resolution states for a job."""
    resolutions = job_manager.get_resolutions(job_id)
    return {"ok": True, "resolutions": resolutions}


@app.post("/api/jobs/{job_id}/resolution")
async def update_resolution(job_id: str, request: Request):
    """Set resolution status and/or append a note for one provision."""
    body = await request.json()
    tenant_idx = body.get("tenant_idx")
    provision_id = body.get("provision_id")
    status = body.get("status")          # optional
    note = body.get("note")              # optional
    notes = body.get("notes")            # optional full note list
    concern_state = body.get("concern_state")
    concern_reason = body.get("concern_reason")

    if tenant_idx is None or not provision_id:
        raise HTTPException(status_code=400, detail="tenant_idx and provision_id required")

    result = job_manager.set_resolution(
        job_id,
        int(tenant_idx),
        provision_id,
        status,
        note,
        notes,
        concern_state,
        concern_reason,
    )
    if result is None:
        raise HTTPException(status_code=404, detail="Job not found")
    return {"ok": True, "resolution": result}



# ══════════════════════════════════════════════════════════════════════════════
# COVERAGE RESOLUTION WORKFLOW
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/jobs/{job_id}/cov-resolutions")
def get_cov_resolutions(job_id: str):
    """Return all coverage resolution states for a job."""
    cov_resolutions = job_manager.get_cov_resolutions(job_id)
    return {"ok": True, "cov_resolutions": cov_resolutions}


@app.post("/api/jobs/{job_id}/cov-resolution")
async def update_cov_resolution(job_id: str, request: Request):
    """Set coverage workflow status for one provision."""
    body = await request.json()
    tenant_idx = body.get("tenant_idx")
    provision_id = body.get("provision_id")
    status = body.get("status", "open")

    if tenant_idx is None or not provision_id:
        raise HTTPException(status_code=400, detail="tenant_idx and provision_id required")

    valid_statuses = ("open", "reviewed", "flagged", "accepted")
    if status not in valid_statuses:
        raise HTTPException(status_code=400, detail=f"status must be one of {valid_statuses}")

    result = job_manager.set_cov_resolution(
        job_id, int(tenant_idx), provision_id, status
    )
    if result is None:
        raise HTTPException(status_code=404, detail="Job not found")
    return {"ok": True, "cov_resolution": result}


# ══════════════════════════════════════════════════════════════════════════════
# PRE-SCAN ENDPOINTS (050)
# ══════════════════════════════════════════════════════════════════════════════

# ── Aggressive Read endpoint (Step 228) ──

@app.post("/api/jobs/{job_id}/aggressive-read")
async def aggressive_read(job_id: str, body: dict):
    """Return the strongest tenant-favorable reading of a specific provision.

    Exploratory only — does not affect scores, governance signals, or pipeline results.
    """
    tenant_index = body.get("tenant_index", 0)
    provision_id = body.get("provision_id", "")

    if not provision_id:
        raise HTTPException(status_code=400, detail="provision_id required")

    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job.get("status") == "completed" and job_manager.is_job_expired(job):
        raise HTTPException(status_code=410, detail="Analysis expired")

    tenants = job.get("input_config", {}).get("tenants", [])
    if tenant_index < 0 or tenant_index >= len(tenants):
        raise HTTPException(status_code=404, detail="Invalid tenant index")

    result_path = tenants[tenant_index].get("result_path")
    if not result_path or not Path(result_path).exists():
        raise HTTPException(status_code=404, detail="Results not available")

    try:
        results = json.loads(Path(result_path).read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as e:
        raise HTTPException(status_code=500, detail=f"Failed to read results: {e}")

    provision = next(
        (p for p in results.get("provisions", []) if p.get("provision_id") == provision_id),
        None,
    )
    if not provision:
        raise HTTPException(status_code=404, detail=f"Provision {provision_id} not found")

    template_text = (provision.get("template_text") or "").strip()
    tenant_text = (provision.get("tenant_text") or "").strip()
    risk_headline = (provision.get("risk_headline") or "").strip()
    provision_name = (provision.get("provision_name") or provision_id).strip()

    if not tenant_text:
        raise HTTPException(status_code=400, detail="No tenant clause text available for this provision")

    system_prompt = (
        "You are a commercial real estate attorney analyzing a lease deviation from the tenant's perspective. "
        "Your job is to identify the strongest possible reading of the tenant's clause — "
        "the interpretation that most benefits the tenant or most burdens the landlord. "
        "Be specific to the actual clause language. Do not generalize or speculate beyond what the text supports."
    )

    user_prompt = f"""PROVISION: {provision_id} — {provision_name}

STANDARD TEMPLATE CLAUSE:
{template_text or '[not present in template — tenant-added clause]'}

TENANT CLAUSE:
{tenant_text}

WHAT CAM IDENTIFIED:
{risk_headline or 'Deviation from standard template'}

Answer these four things specifically:

1. STRONGEST AGGRESSIVE READING: What is the most expansive interpretation of this tenant clause that benefits the tenant or burdens the landlord? Quote the specific language that supports this reading.

2. WHY IT IS PLAUSIBLE: What makes this reading legally defensible? What ambiguity or structure in the clause supports it?

3. LANDLORD PUSHBACK: What is the strongest argument a landlord's attorney would make against this reading? What limits it?

4. BOTTOM LINE: Is this aggressive reading likely to hold up, or is it a stretch? One sentence.

Be direct. Name the specific clause language. Do not hedge unnecessarily."""

    try:
        from cam.core.config import find_and_load_env
        find_and_load_env()
        from cam.core.provider_router import ModelTarget, ProviderRouter, RouterConfig
        target = ModelTarget(
            name="openai:gpt-5.2",
            provider="openai",
            model="gpt-5.2",
            max_output_tokens=800,
            temperature=0.3,
            timeout_sec=60.0,
        )
        router = ProviderRouter([target], RouterConfig())
        adapter = router._get_adapter("openai")
        response = adapter.call(system_prompt, user_prompt, target)
        return {"analysis": response.strip(), "provision_id": provision_id}
    except Exception as e:
        logger.error(f"Aggressive read failed for {provision_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {e}")


# ── Prescan endpoints removed in step 112 (discovery moved into pipeline) ──
# @app.post("/api/prescan/template")
# Prescan replaced by in-pipeline discovery during Stage 2 evaluation.


# @app.post("/api/prescan/tenants")
# Prescan replaced by in-pipeline discovery during Stage 2 evaluation.



# ══════════════════════════════════════════════════════════════════════════════
# TEMPLATE SUMMARY ENDPOINT (Step 138)
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/template/summary")
async def get_template_summary(
    template_file: UploadFile = File(...),
    access_code: str = Form(""),
):
    """Read a template file, run gate check, extract summary fields."""
    from cam.adapters.lease_review.lease_parser import parse_document
    from cam.adapters.lease_review.lease_gate import check_document_is_lease
    from cam.adapters.lease_review.lease_template_reader import read_template_summary
    from cam.core.config import find_and_load_env
    import tempfile

    find_and_load_env()
    config = get_config()

    # Validate access code if configured
    if config["ACCESS_CODE"] and access_code != config["ACCESS_CODE"]:
        raise HTTPException(status_code=401, detail="Invalid access code")

    # Save to temp file for parsing
    suffix = Path(template_file.filename).suffix
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await template_file.read())
        tmp_path = tmp.name

    try:
        text = parse_document(tmp_path)

        # Skip gate check for known demo/sample files
        _demo_names = {"meridian standard template (demo).txt", "template.txt"}
        _is_demo = template_file.filename.lower().strip() in _demo_names
        if _is_demo:
            logger.info("Template gate: skipped (known demo file: %s)", template_file.filename)
        else:
            # Gate check
            gate = check_document_is_lease(text, {})
            if not gate["is_lease"]:
                return {
                    "gate_passed": False,
                    "gate_message": gate["abort_message"],
                    "landlord": "", "property": "",
                    "base_rent": "", "lease_term": "", "governing_law": ""
                }

        # Extract summary
        summary = read_template_summary(text)
        summary["gate_passed"] = True
        summary["gate_message"] = ""
        return summary

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read template: {e}")
    finally:
        os.unlink(tmp_path)


# ══════════════════════════════════════════════════════════════════════════════
# USER RULES ENDPOINTS (Step 140)
# ══════════════════════════════════════════════════════════════════════════════

def _get_user_rules_path() -> Path:
    """Path to the user rules JSON file."""
    data_dir = Path(__file__).parent.parent / "data"
    data_dir.mkdir(exist_ok=True)
    return data_dir / "user_rules.json"


def _verify_access(access_code: str):
    """Verify access code or raise 401."""
    config = get_config()
    if not config["ACCESS_CODE"]:
        return  # No code configured — allow all
    if access_code == config["ACCESS_CODE"]:
        return
    raise HTTPException(status_code=401, detail="Invalid access code")


def _load_user_rules(access_code: str) -> list:
    """Load rules for a given access code. Returns list."""
    code_hash = hashlib.sha256(access_code.encode()).hexdigest()
    path = _get_user_rules_path()
    if not path.exists():
        return []
    try:
        data = json.loads(path.read_text())
        return data.get(code_hash, [])
    except Exception:
        return []


def _save_user_rules(access_code: str, rules: list) -> None:
    """Save rules for a given access code."""
    code_hash = hashlib.sha256(access_code.encode()).hexdigest()
    path = _get_user_rules_path()
    try:
        data = json.loads(path.read_text()) if path.exists() else {}
    except Exception:
        data = {}
    data[code_hash] = rules
    path.write_text(json.dumps(data, indent=2))


@app.get("/api/rules")
async def get_rules(access_code: str = Query(...)):
    """Get all user rules for this access code."""
    _verify_access(access_code)
    rules = _load_user_rules(access_code)
    return {"rules": rules}


@app.post("/api/rules")
async def add_rule(request: Request):
    """Add a new user rule."""
    body = await request.json()
    access_code = body.get("access_code", "")
    rule_text = (body.get("text") or "").strip()
    provision_hint = body.get("provision_hint", "")

    _verify_access(access_code)
    if not rule_text:
        raise HTTPException(status_code=400, detail="Rule text is required")
    if len(rule_text) > 500:
        raise HTTPException(status_code=400, detail="Rule text too long (max 500 chars)")

    rules = _load_user_rules(access_code)
    if len(rules) >= 20:
        raise HTTPException(status_code=400, detail="Maximum 20 rules allowed")

    now = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
    rule_id = f"rule_{now[:10].replace('-','')}_{len(rules)+1:03d}"
    new_rule = {
        "id": rule_id,
        "text": rule_text,
        "provision_hint": provision_hint,
        "created_at": now,
        "enabled": True,
    }
    rules.append(new_rule)
    _save_user_rules(access_code, rules)
    return {"rule": new_rule}


@app.delete("/api/rules/{rule_id}")
async def delete_rule(rule_id: str, access_code: str = Query(...)):
    """Delete a user rule by ID."""
    _verify_access(access_code)
    rules = _load_user_rules(access_code)
    rules = [r for r in rules if r["id"] != rule_id]
    _save_user_rules(access_code, rules)
    return {"ok": True}


# ══════════════════════════════════════════════════════════════════════════════
# AI SUMMARY ENDPOINT (048)
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/ai-summary")
async def ai_summary(request_body: dict):
    """Generate AI summary paragraph for results dashboard."""
    prompt = request_body.get("prompt", "")
    if not prompt:
        return {"summary": ""}

    try:
        from cam.adapters.lease_review.lease_prescan import _call_gemini_sync
        import asyncio
        text = await asyncio.to_thread(_call_gemini_sync, prompt)
        return {"summary": text}
    except Exception as e:
        logger.error(f"AI summary error: {e}")
        return {"summary": ""}


# ══════════════════════════════════════════════════════════════════════════════
# FINAL DRAFT ENDPOINT (144)
# ══════════════════════════════════════════════════════════════════════════════

async def _deprecated_final_draft(request_body: dict):
    """Generate a Final Draft DOCX from per-provision decisions."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    job_id = request_body.get("job_id", "")
    decisions = request_body.get("decisions", {})  # {pid: {choice, text, provision_name, final_verdict, severity}}

    doc = Document()

    # Title
    title = doc.add_heading("Final Draft — Negotiated Lease Terms", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Generated {datetime.now().strftime('%B %d, %Y')}  \u00b7  CAM\u2122 Patent Pending")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph()  # spacer

    # Summary
    kept_template = [pid for pid, d in decisions.items() if d.get("choice") == "template"]
    kept_tenant   = [pid for pid, d in decisions.items() if d.get("choice") == "tenant" and d.get("final_verdict") == "DEVIATES"]
    custom        = [pid for pid, d in decisions.items() if d.get("choice") == "custom"]
    conforms      = [pid for pid, d in decisions.items() if d.get("final_verdict") == "CONFORMS"]

    summary_lines = [
        f"Provisions restored to template: {len(kept_template)}",
        f"Tenant language accepted: {len(kept_tenant)}",
        f"Custom compromise language: {len(custom)}",
        f"Conforming (no change): {len(conforms)}",
    ]
    for line in summary_lines:
        doc.add_paragraph(line)

    doc.add_page_break()

    # ── Per-provision sections ──
    SEVERITY_ORDER = ["CRITICAL", "HIGH", "MEDIUM", "LOW"]

    # Sort: DEVIATES first (by severity), then CONFORMS
    def sort_key(item):
        pid, d = item
        if d.get("final_verdict") == "CONFORMS":
            return (99, pid)
        sev = d.get("severity", "")
        idx = SEVERITY_ORDER.index(sev) if sev in SEVERITY_ORDER else 10
        return (idx, pid)

    sorted_decisions = sorted(decisions.items(), key=sort_key)

    for pid, d in sorted_decisions:
        choice = d.get("choice", "tenant")
        pname = d.get("provision_name", pid)
        verdict = d.get("final_verdict", "")
        severity = d.get("severity", "")
        text = d.get("text", "").strip()

        # Section heading
        heading_text = f"{pid}  {pname}"
        if severity:
            heading_text += f"  \u2014  {severity}"
        doc.add_heading(heading_text, level=2)

        # Decision badge
        badge_map = {
            "template": "\u2713 Restored to standard template",
            "tenant":   "\u2192 Tenant language accepted" if verdict == "DEVIATES" else "\u2713 Conforms \u2014 no change",
            "custom":   "\u26a1 Custom compromise language",
        }
        badge_text = badge_map.get(choice, "")
        badge_para = doc.add_paragraph()
        badge_run = badge_para.add_run(badge_text)
        badge_run.font.size = Pt(9)
        if choice == "template":
            badge_run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
        elif choice == "custom":
            badge_run.font.color.rgb = RGBColor(0x7c, 0x3a, 0xed)
        else:
            badge_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        # Clause text
        if text:
            clause_para = doc.add_paragraph(text)
            clause_para.style = "Normal"
        else:
            doc.add_paragraph("(No clause text recorded)")

        doc.add_paragraph()  # spacer

    # Disclaimer
    doc.add_page_break()
    disc = doc.add_paragraph(
        "Generated by CAM\u2122 \u00b7 Patent Pending \u00b7 "
        "This document is a drafting aid and does not constitute legal advice. "
        "All provisions should be reviewed by qualified legal counsel before execution."
    )
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in disc.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Return as file download
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    from fastapi.responses import StreamingResponse
    filename = f"final_draft_{job_id or 'lease'}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/final-draft")
async def final_draft(request_body: dict):
    """Generate a Final Draft DOCX from per-provision decisions."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    from fastapi.responses import StreamingResponse

    job_id = request_body.get("job_id", "")
    decisions = request_body.get("decisions", {})

    doc = Document()

    # Title
    title = doc.add_heading("Final Draft \u2014 Negotiated Lease Terms", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Generated {datetime.now().strftime('%B %d, %Y')}  \u00b7  CAM\u2122 Patent Pending")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    doc.add_paragraph()

    # Summary counts
    kept_template = [pid for pid, d in decisions.items() if d.get("choice") == "template"]
    kept_tenant   = [pid for pid, d in decisions.items() if d.get("choice") == "tenant" and d.get("final_verdict") == "DEVIATES"]
    custom        = [pid for pid, d in decisions.items() if d.get("choice") == "custom"]
    conforms      = [pid for pid, d in decisions.items() if d.get("final_verdict") == "CONFORMS"]

    for line in [
        f"Provisions restored to template: {len(kept_template)}",
        f"Tenant language accepted: {len(kept_tenant)}",
        f"Custom compromise language: {len(custom)}",
        f"Conforming (no change): {len(conforms)}",
    ]:
        doc.add_paragraph(line)

    doc.add_page_break()

    SEVERITY_ORDER = ["CRITICAL", "HIGH", "MEDIUM", "LOW"]

    def sort_key(item):
        pid, d = item
        if d.get("final_verdict") == "CONFORMS":
            return (99, pid)
        sev = d.get("severity", "")
        idx = SEVERITY_ORDER.index(sev) if sev in SEVERITY_ORDER else 10
        return (idx, pid)

    for pid, d in sorted(decisions.items(), key=sort_key):
        choice = d.get("choice", "tenant")
        pname = d.get("provision_name", pid)
        verdict = d.get("final_verdict", "")
        severity = d.get("severity", "")
        text = d.get("text", "").strip()

        heading_text = f"{pid}  {pname}"
        if severity:
            heading_text += f"  \u2014  {severity}"
        doc.add_heading(heading_text, level=2)

        badge_map = {
            "template": "\u2713 Restored to standard template",
            "tenant":   "\u2192 Tenant language accepted" if verdict == "DEVIATES" else "\u2713 Conforms \u2014 no change",
            "custom":   "\u26a1 Custom compromise language",
        }
        badge_para = doc.add_paragraph()
        badge_run = badge_para.add_run(badge_map.get(choice, ""))
        badge_run.font.size = Pt(9)
        badge_run.font.color.rgb = (
            RGBColor(0x16, 0xa3, 0x4a) if choice == "template" else
            RGBColor(0x7c, 0x3a, 0xed) if choice == "custom" else
            RGBColor(0x64, 0x74, 0x8b)
        )

        doc.add_paragraph(text if text else "(No clause text recorded)")
        doc.add_paragraph()

    doc.add_page_break()
    disc = doc.add_paragraph(
        "Generated by CAM\u2122 \u00b7 Patent Pending \u00b7 "
        "This document is a drafting aid and does not constitute legal advice. "
        "All provisions should be reviewed by qualified legal counsel before execution."
    )
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in disc.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"final_draft_{job_id or 'lease'}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ══════════════════════════════════════════════════════════════════════════════
# LEASE-SPECIFIC ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/jobs/lease")
async def create_lease_job(
    access_code: str = Form(...),
    email: Optional[str] = Form(None),
    template_file: Optional[UploadFile] = File(None),
    tenant_files: List[UploadFile] = File(...),
    provisions: Optional[str] = Form(None),
    custom_provisions: Optional[str] = Form(None),
    custom_from_scan: Optional[str] = Form(None),
    added_from_scan: Optional[str] = Form(None),
    prescan_record: Optional[str] = Form(None),
    instructions: Optional[str] = Form(None),
    strictness: Optional[str] = Form("standard"),
    template_type: Optional[str] = Form("blank_template"),
    identity_check: Optional[str] = Form("landlord_property"),
    mode: Optional[str] = Form("compare"),
    # Step 261: perspective (tenant / landlord / neutral) is required — the
    # frontend gates submit on this. Backend validates and persists; prompts
    # don't yet branch on it (Step 262 work).
    perspective: Optional[str] = Form(None),
):
    """Create a lease analysis job with file uploads."""
    config = get_config()

    # Validate access code
    if config["ACCESS_CODE"] and access_code != config["ACCESS_CODE"]:
        raise HTTPException(status_code=401, detail="Invalid access code")

    # Validate mode
    if mode not in ("compare", "analyze"):
        raise HTTPException(status_code=400, detail="mode must be 'compare' or 'analyze'")

    # Step 261: validate perspective. Required — frontend gates on this, but the
    # API enforces it independently in case a non-browser caller hits the endpoint.
    if perspective not in ("tenant", "landlord", "neutral"):
        raise HTTPException(
            status_code=400,
            detail="perspective is required and must be one of 'tenant', 'landlord', 'neutral'",
        )

    # Mode A (compare) requires a template; Mode C (analyze) does not.
    if mode == "compare" and template_file is None:
        raise HTTPException(status_code=400, detail="template_file is required for mode='compare'")

    # Validate strictness
    if strictness not in ("permissive", "standard", "strict"):
        raise HTTPException(status_code=400, detail="strictness must be permissive, standard, or strict")

    # Validate template_type
    if template_type not in ("blank_template", "executed_reference"):
        template_type = "blank_template"

    # Validate identity_check
    if identity_check not in ("clauses_only", "landlord_property", "landlord_tenant"):
        identity_check = "landlord_property"

    # Parse provisions JSON (if provided)
    selected_ids = None
    if provisions:
        try:
            selected_ids = json.loads(provisions)
            if not isinstance(selected_ids, list):
                raise ValueError("provisions must be a JSON array")
        except (json.JSONDecodeError, ValueError) as e:
            raise HTTPException(status_code=400, detail=f"Invalid provisions JSON: {e}")

    # Parse custom provisions JSON (if provided)
    custom_provisions_list = None
    if custom_provisions:
        try:
            custom_provisions_list = json.loads(custom_provisions)
            if not isinstance(custom_provisions_list, list):
                raise ValueError("custom_provisions must be a JSON array")
        except (json.JSONDecodeError, ValueError) as e:
            raise HTTPException(status_code=400, detail=f"Invalid custom_provisions JSON: {e}")

    # Parse pre-scan provisions (050)
    custom_from_scan_list = []
    if custom_from_scan:
        try:
            custom_from_scan_list = json.loads(custom_from_scan)
        except (json.JSONDecodeError, ValueError):
            custom_from_scan_list = []

    # Parse prescan record (092)
    prescan_record_parsed = None
    if prescan_record:
        try:
            prescan_record_parsed = json.loads(prescan_record)
        except (json.JSONDecodeError, ValueError):
            prescan_record_parsed = None

    added_from_scan_list = []
    if added_from_scan:
        try:
            added_from_scan_list = json.loads(added_from_scan)
        except (json.JSONDecodeError, ValueError):
            added_from_scan_list = []

    # Generate job ID early so we can use it for upload directory
    job_id = job_manager._generate_job_id("lease_review")

    # Save uploaded files
    upload_dir = Path(config["UPLOAD_DIR"]) / job_id
    upload_dir.mkdir(parents=True, exist_ok=True)

    # Save template (Mode A only; Mode C has no template)
    template_save_path = None
    if template_file is not None:
        template_save_path = upload_dir / template_file.filename
        content = await template_file.read()
        template_save_path.write_bytes(content)

    # Save tenant files and build tenant list (with ZIP extraction)
    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
    tenants = []
    for tf in tenant_files:
        tenant_save_path = upload_dir / tf.filename
        content = await tf.read()
        tenant_save_path.write_bytes(content)

        if tenant_save_path.suffix.lower() == ".zip":
            # Extract supported files from ZIP
            try:
                with zipfile.ZipFile(str(tenant_save_path), "r") as zf:
                    for name in zf.namelist():
                        # Skip directories and hidden files
                        if name.endswith("/") or name.startswith("__") or "/." in name:
                            continue
                        ext = Path(name).suffix.lower()
                        if ext not in ALLOWED_EXTENSIONS:
                            continue
                        # Extract to upload dir
                        extracted_name = Path(name).name  # strip subdirectories
                        extracted_path = upload_dir / extracted_name
                        with zf.open(name) as src:
                            extracted_path.write_bytes(src.read())
                        tenants.append({
                            "filename": extracted_name,
                            "upload_path": str(extracted_path),
                            "status": "queued",
                            "stage": None,
                            "error": None,
                            "result_path": None,
                            "annotated_path": None,
                        })
                # Remove the ZIP after extraction
                tenant_save_path.unlink(missing_ok=True)
            except zipfile.BadZipFile:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid ZIP file: {tf.filename}",
                )
        else:
            tenants.append({
                "filename": tf.filename,
                "upload_path": str(tenant_save_path),
                "status": "queued",
                "stage": None,
                "error": None,
                "result_path": None,
                "annotated_path": None,
            })

    if not tenants:
        raise HTTPException(
            status_code=400,
            detail="No valid tenant files found. Supported formats: PDF, DOCX, TXT (or ZIP containing these).",
        )

    # Build input config
    input_config = {
        "template_file": template_file.filename if template_file is not None else "",
        "template_path": str(template_save_path) if template_save_path is not None else "",
        "tenants": tenants,
        "provisions": selected_ids,
        "custom_provisions": custom_provisions_list,
        "custom_from_scan": custom_from_scan_list,
        "added_from_scan": added_from_scan_list,
        "prescan_record": prescan_record_parsed,
        "strictness": strictness,
        "instructions": instructions or "",
        "template_type": template_type,
        "identity_check": identity_check,
        "access_code": access_code,  # Step 140: for user rules injection
        "mode": mode,
        # Step 261: perspective lens (tenant / landlord / neutral). Persisted now;
        # downstream prompts will branch on it in Step 262.
        "perspective": perspective,
    }

    # Create job with the pre-generated ID (matches upload directory)
    job = job_manager.create_job("lease_review", email, input_config, job_id=job_id)

    # Start background processing
    job_manager.start_processing(job_id)
    started_job = job_manager.get_job(job_id) or job

    return {
        "job_id": job_id,
        "status": "processing",
        "estimated_minutes": job["estimated_minutes"],
        "started_at": started_job.get("started_at"),
        "created_at": started_job.get("created_at"),
        "results_url": f"{config['APP_BASE_URL']}/results/{job_id}",
    }


@app.get("/api/jobs/{job_id}/results/{tenant_index}/annotated")
def download_annotated(job_id: str, tenant_index: int, with_resolutions: bool = True):
    """Download the annotated document (DOCX or PDF) for a tenant.

    If with_resolutions=True (default), regenerates the DOCX on-demand so it
    includes the lawyer's latest CAM review decisions and notes.
    PDF annotated files are always served as-is (no regeneration).
    """
    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job.get("status") == "completed" and job_manager.is_job_expired(job):
        return JSONResponse(status_code=410, content={"detail": "This analysis has expired."})

    tenants = job.get("input_config", {}).get("tenants", [])
    if tenant_index < 0 or tenant_index >= len(tenants):
        raise HTTPException(status_code=404, detail="Invalid tenant index")

    tenant = tenants[tenant_index]
    annotated_path = tenant.get("annotated_path")

    if not annotated_path or not Path(annotated_path).exists():
        raise HTTPException(status_code=404, detail="Annotated document not available")

    ext = Path(annotated_path).suffix.lower()
    filename = f"annotated_{tenant['filename']}"
    if not filename.lower().endswith(ext):
        filename = Path(filename).stem + ext

    # For DOCX or PDF: regenerate on-demand with latest resolutions
    if with_resolutions and ext in (".docx", ".pdf"):
        result_path = tenant.get("result_path")
        upload_path = tenant.get("upload_path")
        resolutions = job.get("resolutions", {})

        if result_path and Path(result_path).exists() and upload_path and Path(upload_path).exists():
            try:
                pipeline_results = json.loads(Path(result_path).read_text(encoding="utf-8"))

                if ext == ".docx":
                    from cam.adapters.lease_review.lease_docx_annotator import annotate_docx
                    cov_resolutions = job.get("cov_resolutions", {})
                    regen_path = Path(annotated_path).parent / f"annotated_latest_{Path(tenant['filename']).stem}.docx"
                    annotate_docx(
                        original_docx_path=upload_path,
                        results=pipeline_results,
                        output_path=str(regen_path),
                        resolutions=resolutions,
                        cov_resolutions=cov_resolutions,
                    )
                    return FileResponse(
                        path=str(regen_path),
                        filename=filename,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                else:  # .pdf
                    from cam.adapters.lease_review.lease_pdf_annotator import annotate_pdf
                    cov_resolutions = job.get("cov_resolutions", {})
                    regen_path = Path(annotated_path).parent / f"annotated_latest_{Path(tenant['filename']).stem}.pdf"
                    annotate_pdf(
                        original_pdf_path=upload_path,
                        results=pipeline_results,
                        output_path=str(regen_path),
                        resolutions=resolutions,
                        cov_resolutions=cov_resolutions,
                    )
                    return FileResponse(
                        path=str(regen_path),
                        filename=filename,
                        media_type="application/pdf",
                    )
            except Exception as e:
                logger.warning(f"Annotated document regeneration failed, falling back to cached: {e}")
                # Fall through to serve cached file

    return FileResponse(
        path=annotated_path,
        filename=filename,
        media_type="application/octet-stream",
    )


@app.get("/api/jobs/{job_id}/results/{tenant_index}/comparison")
def download_comparison_view(job_id: str, tenant_index: int):
    """Step 255: download the Aligned Provision Comparison View PDF for a tenant.

    Mode A only — Mode C jobs do not generate this artifact, so the path will
    be absent on the tenant record. Generation already happened during the
    pipeline run (additive call inside `generate_outputs`); this endpoint just
    serves the existing file.
    """
    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job.get("status") == "completed" and job_manager.is_job_expired(job):
        return JSONResponse(status_code=410, content={"detail": "This analysis has expired."})

    tenants = job.get("input_config", {}).get("tenants", [])
    if tenant_index < 0 or tenant_index >= len(tenants):
        raise HTTPException(status_code=404, detail="Invalid tenant index")

    tenant = tenants[tenant_index]
    comparison_path = tenant.get("comparison_view_path")

    if not comparison_path or not Path(comparison_path).exists():
        raise HTTPException(
            status_code=404,
            detail="Aligned Provision Comparison not available for this tenant",
        )

    tenant_stem = Path(tenant.get("filename", "lease")).stem or "lease"
    filename = f"{tenant_stem}_Aligned_Provision_Comparison.pdf"
    return FileResponse(
        path=comparison_path,
        filename=filename,
        media_type="application/pdf",
    )


@app.post("/api/jobs/{job_id}/add-tenants")
async def add_tenants_to_job(
    job_id: str,
    tenant_files: List[UploadFile] = File(...),
    add_provisions: Optional[str] = Form(None),
):
    """Add more tenant leases to an existing completed job.

    Reuses the cached standard template from the original analysis.
    New tenants are processed and appended to the existing job's results.
    """
    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job.get("status") not in ("completed", "cancelled"):
        raise HTTPException(status_code=400, detail="Job must be completed before adding tenants")

    if job_manager.is_job_expired(job):
        return JSONResponse(status_code=410, content={"detail": "This analysis has expired."})

    config = get_config()
    input_cfg = job["input_config"]
    template_path = input_cfg.get("template_path")

    if not template_path or not Path(template_path).exists():
        # Template was already deleted — need to re-upload
        raise HTTPException(
            status_code=400,
            detail="Template file no longer available. Please start a new analysis.",
        )

    # Save new tenant files
    upload_dir = Path(config["UPLOAD_DIR"]) / job_id
    upload_dir.mkdir(parents=True, exist_ok=True)

    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
    new_tenants = []
    for tf in tenant_files:
        tenant_save_path = upload_dir / tf.filename
        content = await tf.read()
        tenant_save_path.write_bytes(content)

        if tenant_save_path.suffix.lower() == ".zip":
            try:
                with zipfile.ZipFile(str(tenant_save_path), "r") as zf:
                    for name in zf.namelist():
                        if name.endswith("/") or name.startswith("__") or "/." in name:
                            continue
                        ext = Path(name).suffix.lower()
                        if ext not in ALLOWED_EXTENSIONS:
                            continue
                        extracted_name = Path(name).name
                        extracted_path = upload_dir / extracted_name
                        with zf.open(name) as src:
                            extracted_path.write_bytes(src.read())
                        new_tenants.append({
                            "filename": extracted_name,
                            "upload_path": str(extracted_path),
                            "status": "queued",
                            "stage": None,
                            "error": None,
                            "result_path": None,
                            "annotated_path": None,
                        })
                tenant_save_path.unlink(missing_ok=True)
            except zipfile.BadZipFile:
                raise HTTPException(status_code=400, detail=f"Invalid ZIP file: {tf.filename}")
        else:
            new_tenants.append({
                "filename": tf.filename,
                "upload_path": str(tenant_save_path),
                "status": "queued",
                "stage": None,
                "error": None,
                "result_path": None,
                "annotated_path": None,
            })

    if not new_tenants:
        raise HTTPException(status_code=400, detail="No valid tenant files found.")

    # If caller also wants new provisions added in the same pass, apply them first
    # so that incremental processing picks them up for the new tenants.
    # (Existing tenants do NOT get re-run here — user can do a separate add-provisions
    #  run for existing tenants once this completes.)
    new_provision_ids = []
    if add_provisions:
        try:
            parsed = json.loads(add_provisions)
            if isinstance(parsed, list):
                existing_ids = set(input_cfg.get("provisions") or [])
                new_provision_ids = [p for p in parsed if p not in existing_ids]
                if new_provision_ids:
                    job_manager.append_provisions(job_id, new_provision_ids)
        except (json.JSONDecodeError, TypeError):
            pass  # Ignore malformed field — proceed with tenants only

    # Append new tenants to the job and set back to processing
    job_manager.append_tenants(job_id, new_tenants)

    # Start background processing for new tenants only
    job_manager.start_incremental_processing(job_id, len(input_cfg["tenants"]) - len(new_tenants))

    return {
        "job_id": job_id,
        "status": "processing",
        "started_at": (job_manager.get_job(job_id) or {}).get("started_at"),
        "new_tenants": len(new_tenants),
        "new_provisions_added": new_provision_ids,
        "total_tenants": len(input_cfg["tenants"]),
    }


@app.post("/api/jobs/{job_id}/add-provisions")
async def add_provisions_to_job(
    job_id: str,
    body: dict,
):
    """Add more provisions to an existing completed job.

    Runs the analysis pipeline on new provisions only for each tenant,
    using cached tenant text and standard provisions.
    """
    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job.get("status") not in ("completed", "cancelled"):
        raise HTTPException(status_code=400, detail="Job must be completed before adding provisions")

    if job_manager.is_job_expired(job):
        return JSONResponse(status_code=410, content={"detail": "This analysis has expired."})

    new_provision_ids = body.get("provisions", [])
    if not new_provision_ids or not isinstance(new_provision_ids, list):
        raise HTTPException(status_code=400, detail="Required: provisions (list of provision IDs)")

    config = get_config()
    input_cfg = job["input_config"]
    template_path = input_cfg.get("template_path")

    if not template_path or not Path(template_path).exists():
        raise HTTPException(
            status_code=400,
            detail="Template file no longer available. Please start a new analysis.",
        )

    # Get existing provision IDs to avoid re-running
    existing_ids = set(input_cfg.get("provisions") or [])
    truly_new = [pid for pid in new_provision_ids if pid not in existing_ids]

    if not truly_new:
        raise HTTPException(status_code=400, detail="All requested provisions were already analyzed.")

    # Update job's provision list and set back to processing
    job_manager.append_provisions(job_id, truly_new)

    # Start background processing for new provisions
    job_manager.start_provision_processing(job_id, truly_new)

    return {
        "job_id": job_id,
        "status": "processing",
        "started_at": (job_manager.get_job(job_id) or {}).get("started_at"),
        "new_provisions": truly_new,
        "total_provisions": len(existing_ids | set(truly_new)),
    }


@app.post("/api/jobs/{job_id}/chat")
async def chat_with_analysis(job_id: str, body: dict):
    """Ask follow-up questions about the analysis results.

    Supports single-model (default) and multi-model modes.
    Full analysis context is included in the system prompt.
    """
    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    # Allow chat on completed, cancelled, or processing jobs (partial results are fine —
    # the user may be reviewing a finished tenant while others are still running)
    if job.get("status") not in ("completed", "cancelled", "processing"):
        raise HTTPException(status_code=400, detail="Job must be completed to chat")

    if job_manager.is_job_expired(job):
        return JSONResponse(status_code=410, content={"detail": "This analysis has expired."})

    question = body.get("question", "").strip()
    if not question:
        raise HTTPException(status_code=400, detail="Question is required")

    mode = body.get("mode", "single")  # "single" or "multi"
    models = body.get("models", ["claude"])
    synthesize = body.get("synthesize", False)
    synthesizer = body.get("synthesizer", "claude")
    provision_id = body.get("provision_id")  # optional scope
    tenant_idx = body.get("tenant_idx")  # optional int — scope to one tenant
    ui_context = body.get("ui_context", {}) or {}
    history = body.get("history", [])

    # Load tenant results for context
    domain_label = "Document Analysis"  # safe default
    # Step 260: detect Mode C (single-doc coverage analysis) so we can build a
    # coverage-shaped context and prompt instead of the deviation-shaped Mode A
    # version. The mode field is set by the upload endpoint (Step 252).
    is_mode_c = (job.get("input_config", {}) or {}).get("mode") == "analyze"
    # Step 262: read perspective so the SCOPE block can frame coverage gaps
    # and risk against the selected party (tenant / landlord / neutral).
    # Older jobs without this field default to tenant for back-compat.
    job_perspective = ((job.get("input_config", {}) or {}).get("perspective") or "tenant").lower()
    if job_perspective not in ("tenant", "landlord", "neutral"):
        job_perspective = "tenant"
    tenants = job.get("input_config", {}).get("tenants", [])
    context_parts = []
    for i, tenant in enumerate(tenants):
        # If tenant_idx specified, only include that tenant
        if tenant_idx is not None and i != int(tenant_idx):
            continue
        result_path = tenant.get("result_path")
        if result_path and Path(result_path).exists():
            try:
                data = json.loads(Path(result_path).read_text(encoding="utf-8"))
                # Read domain label from first result (pipeline-agnostic)
                if i == 0:
                    domain_label = data.get("pipeline_domain_label", domain_label)
                tenant_file = data.get("tenant_file", tenant.get("filename", ""))
                context_parts.append(f"\n--- Tenant: {tenant_file} ---")

                if provision_id:
                    # Scope to specific provision (Mode C: match by issue_area_id)
                    if is_mode_c:
                        for a in data.get("coverage_assessment", []):
                            if a.get("issue_area_id") == provision_id:
                                context_parts.append(json.dumps(a, indent=2, default=str))
                                break
                    else:
                        for p in data.get("provisions", []):
                            if p.get("provision_id") == provision_id:
                                context_parts.append(json.dumps(p, indent=2, default=str))
                                break
                elif is_mode_c:
                    # Mode C: build summary from coverage_assessment instead of provisions.
                    # Each item carries coverage_state (covered / partial / missing /
                    # covered_unfavorable / not_applicable), an optional partial_class
                    # (partial_material vs partial_review), an exposure_statement (why this
                    # gap matters in lawyer-facing language), and optional elements_missing
                    # / negative_space_signals lists.
                    ca_list = data.get("coverage_assessment", []) or []
                    for a in ca_list:
                        iaid = a.get("issue_area_id", "")
                        iname = a.get("issue_area_name", iaid)
                        state = a.get("coverage_state", "")
                        pclass = a.get("partial_class", "")
                        stmt = (a.get("exposure_statement", "") or "").strip()
                        missing = a.get("elements_missing", []) or []
                        ns_signals = a.get("negative_space_signals", []) or []
                        head = f"  {iaid}: {iname} — coverage_state={state}"
                        if pclass:
                            head += f", partial_class={pclass}"
                        body_lines = [head]
                        if stmt:
                            body_lines.append(f"    Exposure: {stmt}")
                        if missing:
                            body_lines.append(f"    Missing elements: {', '.join(str(m) for m in missing[:5])}")
                        if ns_signals:
                            ns_summary = ", ".join(
                                str(s.get("signal_type", "") or s.get("description", ""))
                                for s in ns_signals[:3]
                            )
                            if ns_summary:
                                body_lines.append(f"    Negative space: {ns_summary}")
                        context_parts.append("\n".join(body_lines))
                    meta = data.get("contract_metadata", {})
                    if meta:
                        context_parts.append(f"  Contract: {json.dumps(meta, default=str)}")
                    # Step 312: include Stage 7 cross-provision findings
                    cpfs = data.get("cross_provision_findings") or []
                    if cpfs:
                        context_parts.append(f"  Stage 7 — Contract Interaction Review ({len(cpfs)} finding(s)):")
                        for cpf in cpfs:
                            ftype = cpf.get("finding_type", "")
                            fid   = cpf.get("finding_id", "")
                            lps   = ", ".join(cpf.get("implicated_lps") or [])
                            hl    = (cpf.get("headline") or "").strip()
                            sev   = cpf.get("severity", "")
                            agree = cpf.get("evaluator_agreement", "")
                            direc = cpf.get("directionality") or ""
                            context_parts.append(
                                f"    {fid} [{ftype}] {lps} — {hl}"
                                + (f" | severity={sev}" if sev else "")
                                + (f" | agreement={agree}" if agree else "")
                                + (f" | directionality={direc}" if direc else "")
                            )
                else:
                    # Include all provisions summary
                    for p in data.get("provisions", []):
                        pid = p.get("provision_id", "")
                        verdict = p.get("final_verdict", "")
                        sev = p.get("severity", "")
                        headline = p.get("risk_headline", "")
                        details = p.get("challenge_details", "")
                        action = p.get("recommended_action", "")
                        agreement = p.get("agreement_pattern", "")
                        ev_verdicts = p.get("evaluator_verdicts", {})
                        ev_summary = ", ".join(f"{k}: {v}" for k, v in ev_verdicts.items()) if ev_verdicts else ""
                        context_parts.append(
                            f"  {pid}: {verdict} ({sev}) - {headline}\n"
                            f"    Evaluator agreement: {agreement}{(' (' + ev_summary + ')') if ev_summary else ''}\n"
                            f"    Details: {details}\n"
                            f"    Action: {action}"
                        )

                    # Include key text excerpts
                    meta = data.get("contract_metadata", {})
                    if meta:
                        context_parts.append(f"  Contract: {json.dumps(meta, default=str)}")
            except (json.JSONDecodeError, OSError):
                pass

    analysis_context = "\n".join(context_parts)
    ui_context_lines = []
    if ui_context:
        screen = ui_context.get("screen")
        if screen:
            ui_context_lines.append(f"Current screen: {screen}")
        if ui_context.get("mode"):
            ui_context_lines.append(f"Mode: {ui_context.get('mode')}")
        if "mode_explicitly_selected" in ui_context:
            ui_context_lines.append(
                f"Mode explicitly selected by user: {'yes' if ui_context.get('mode_explicitly_selected') else 'no'}"
            )
        if ui_context.get("perspective"):
            ui_context_lines.append(f"Perspective: {ui_context.get('perspective')}")
        top_tab = ui_context.get("active_top_tab") or {}
        if top_tab.get("label"):
            ui_context_lines.append(f"Active top tab: {top_tab.get('label')} ({top_tab.get('id', '')})")
        result_tab = ui_context.get("active_results_tab") or {}
        if result_tab.get("label"):
            ui_context_lines.append(f"Active detail tab: {result_tab.get('label')} ({result_tab.get('id', '')})")
        if "contract_detail_open" in ui_context:
            ui_context_lines.append(f"Contract detail open: {'yes' if ui_context.get('contract_detail_open') else 'no'}")
        current_contract = ui_context.get("current_contract") or {}
        if current_contract.get("label"):
            ui_context_lines.append(f"Current contract: {current_contract.get('label')}")
        chat_scope = ui_context.get("chat_scope") or {}
        if chat_scope.get("contract_label"):
            ui_context_lines.append(f"Chat scope contract: {chat_scope.get('contract_label')}")
        if chat_scope.get("provision_label"):
            ui_context_lines.append(f"Chat scope provision: {chat_scope.get('provision_label')}")
        views = ui_context.get("available_views") or []
        if views:
            ui_context_lines.append("Available views:")
            for view in views:
                label = view.get("label") or view.get("id") or "View"
                purpose = view.get("purpose") or ""
                ui_context_lines.append(f"- {label}: {purpose}".rstrip(": "))
    ui_context_block = "\n".join(ui_context_lines)
    ui_context_prompt = f"UI CONTEXT:\n{ui_context_block}\n\n" if ui_context_block else ""

    # Step 260: Mode C uses a coverage-oriented SCOPE block; Mode A keeps the
    # deviation-oriented one. Built as a separate variable to keep the prompt
    # concatenation readable.
    if is_mode_c:
        scope_block = (
            "SCOPE (Mode C — single-document coverage analysis):\n"
            "There is no reference template for this run. Findings are framed as coverage of an\n"
            "issue-area schema (typically 18 commercial-lease issue areas). Each issue area has:\n"
            "  - coverage_state: 'covered' (clause is present and favorable),\n"
            "    'covered_unfavorable' (clause is present but problematic for the tenant),\n"
            "    'partial' (clause is present but incomplete),\n"
            "    'missing' (clause is absent),\n"
            "    'not_applicable' (issue does not apply to this lease type).\n"
            "  - partial_class (only when state=partial): 'partial_material' = needs attention,\n"
            "    'partial_review' = worth reviewing.\n"
            "  - exposure_statement: lawyer-facing explanation of why the gap matters.\n"
            "  - elements_missing: bullet list of specific clause elements that are absent.\n"
            "  - negative_space_signals: structural absence cues (broken cross-references,\n"
            "    missing exhibits, reserved sections).\n\n"
            "User-facing buckets the UI uses (mirror these in your answers):\n"
            "  - 'need attention' = covered_unfavorable, missing, or partial_material\n"
            "  - 'worth reviewing' = partial_review\n"
            "  - 'covered' = covered\n"
            "  - 'not applicable' = not_applicable\n\n"
            "Answer questions about which issue areas have coverage gaps, what language the user\n"
            "should push for, what the risk is of leaving a gap unaddressed, and how to negotiate\n"
            "missing or unfavorable clauses. Do NOT use Mode A vocabulary like 'deviation',\n"
            "'flagged', 'severity (Critical/High/Medium/Low)', 'evaluator agreement', or 'reference\n"
            "language' — those concepts don't apply here. When asked 'what should I do', anchor\n"
            "the answer in the actual coverage_state and exposure_statement of the issue areas.\n\n"
            + _build_perspective_addendum(job_perspective)
        )
    else:
        scope_block = (
            "SCOPE:\n"
            "Answer questions about the documents, findings, provisions, deviations, risk implications, "
            "negotiation strategy, and legal concepts relevant to this analysis. "
            + _build_perspective_addendum(job_perspective)
        )

    system_prompt = (
        f"You are a specialized assistant for {domain_label}. "
        f"You have access to the structured analysis findings for the document(s) in this session, "
        f"provided below as ANALYSIS CONTEXT.\n\n"

        + scope_block +
        "If asked whether to sign, proceed, or how to approach a negotiation, engage fully — "
        "use the actual findings to give a substantive answer. "
        "You are not a licensed attorney and this is not legal advice; say so naturally at the end "
        "when your answer touches on decisions or professional judgment, not as an upfront disclaimer. "
        "If a question is clearly unrelated to this analysis or its subject matter, redirect warmly: "
        f"'I'm focused on your {domain_label.lower()} — happy to dig into any of the findings "
        "or documents in this session. Is there something specific you would like to explore?'\n\n"

        "RESPONSE GUIDELINES:\n"
        "- Be specific: reference actual provisions, findings, and severity levels from the analysis.\n"
        "- Plain language. This is a conversation, not a memo.\n"
        "- 200-300 words unless the user asks for more detail.\n"
        "- No markdown code blocks or code fences. Use quotation marks for contract language.\n"
        "- In multi-model responses, refer to models by name (Claude, GPT-5.2, Grok, Gemini).\n"
        "- When disclaiming, do it once, briefly, at the end — not at the start, not repeatedly.\n\n"

        "- If UI CONTEXT is provided, use the exact visible screen and tab labels from it.\n"
        "- Only reference views, tabs, or interface elements that appear in UI CONTEXT.\n"
        "- When the user asks where to find something, answer using UI CONTEXT first.\n"
        "- If the user is already on the correct screen or tab, say so directly.\n\n"

        f"{ui_context_prompt}"
        f"CAM KNOWLEDGE BASE (use this to answer any questions about CAM, scores, signals, pipeline stages, or provisions):\n{CAM_KNOWLEDGE}\n\n"

        f"ANALYSIS CONTEXT:\n{analysis_context}"
    )

    # Build conversation context from history (Part 4: memory fix)
    conversation_context = ""
    if history:
        recent = history[-10:]  # Last 10 turns (5 exchanges)
        for h in recent:
            role = h.get("role", "user")
            content = h.get("content", "")
            if role == "user":
                conversation_context += f"\nUser asked: {content}\n"
            else:
                # Check if this is a multi-model response (contains [Model]: labels)
                if content.startswith("[") and "]:" in content:
                    conversation_context += f"\nMultiple AI models responded:\n{content}\n"
                else:
                    conversation_context += f"\nAI responded: {content}\n"

    # Build the full user prompt with history
    if conversation_context:
        full_prompt = (
            f"CONVERSATION HISTORY:\n{conversation_context}\n"
            f"---\n"
            f"CURRENT QUESTION:\n{question}"
        )
    else:
        full_prompt = question

    # Route to model(s)
    model_map = {
        "claude": "claude",
        "gpt": "openai",
        "grok": "xai",
        "gemini": "google",
    }
    model_labels = {"claude": "Claude", "openai": "GPT-5.2", "xai": "Grok", "google": "Gemini"}

    from cam.adapters.lease_review.model_config import get_display_name as _get_display_name

    def _friendly_model_label(model_str: str, fallback_key: str) -> str:
        return _get_display_name(model_str, model_labels.get(fallback_key, fallback_key))

    try:
        from cam.core.llm import call_llm
        import concurrent.futures

        if mode == "multi":
            # Multi-model: parallel calls
            responses = {}
            actual_models = {}
            selected_models = [m for m in models if m in model_map]
            if not selected_models:
                selected_models = ["claude"]

            def _call_model(model_key):
                provider = model_map[model_key]
                try:
                    result = call_llm(
                        provider=provider,
                        system_prompt=system_prompt,
                        user_prompt=full_prompt,
                        temperature=0.3,
                    )
                    actual_model = result.get("model", "")
                    return model_key, result.get("content", ""), actual_model
                except Exception as e:
                    return model_key, f"Error: {e}", ""

            with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
                futures = {executor.submit(_call_model, m): m for m in selected_models}
                for future in concurrent.futures.as_completed(futures):
                    model_key, response, actual_model = future.result()
                    responses[model_key] = response
                    actual_models[model_key] = actual_model

            if synthesize and len(responses) > 1:
                # Build synthesis prompt with all model responses
                model_responses_text = "\n\n".join(
                    f"[{k.upper()}]: {v}" for k, v in responses.items()
                )
                synthesis_prompt = (
                    f"The following question was asked:\n{question}\n\n"
                    f"These AI models responded:\n\n{model_responses_text}\n\n"
                    f"Synthesize these responses into one clear, unified answer. "
                    f"Note where models agreed and where they differed. "
                    f"Attribute specific insights to the model that provided them. "
                    f"Be concise (200-300 words)."
                )
                synthesizer_provider = model_map.get(synthesizer, "anthropic")
                model_labels_display = {"claude": "Claude", "gpt": "GPT-5.2", "grok": "Grok", "gemini": "Gemini"}
                try:
                    synthesis_result = _call_chat_with_followups(
                        call_llm=call_llm,
                        provider=synthesizer_provider,
                        system_prompt=system_prompt,
                        user_prompt=synthesis_prompt,
                        temperature=0.3,
                    )
                    return {
                        "synthesized_response": synthesis_result.get("content", ""),
                        "suggested_followups": synthesis_result.get("suggested_followups", []),
                        "synthesized_by": model_labels_display.get(synthesizer, synthesizer),
                        "mode": "multi",
                        "individual_responses": responses,
                        "actual_models": actual_models,
                    }
                except Exception as e:
                    # Fall back to individual if synthesis fails
                    logger.warning(f"Synthesis failed, falling back to individual: {e}")
                    return {"responses": responses, "actual_models": actual_models, "mode": "multi"}
            else:
                return {"responses": responses, "actual_models": actual_models, "mode": "multi"}

        else:
            # Single model — use selected model from picker (default: Claude)
            selected_model = body.get("model", "claude")
            provider = model_map.get(selected_model, "claude")
            result = _call_chat_with_followups(
                call_llm=call_llm,
                provider=provider,
                system_prompt=system_prompt,
                user_prompt=full_prompt,
                temperature=0.3,
            )

            return {
                "response": result.get("content", ""),
                "suggested_followups": result.get("suggested_followups", []),
                "mode": "single",
                "model_label": _friendly_model_label(result.get("model", ""), selected_model),
                "model_key": selected_model,
                "actual_model": result.get("model", ""),
            }

    except Exception as e:
        logger.error(f"Chat error for job {job_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Chat error: {e}")


@app.post("/api/chat/general")
async def chat_general(request: Request):
    """General lease guidance chat — no job required.

    For pre-analysis questions about review areas, lease concepts,
    and what to look for.
    """
    body = await request.json()
    question = body.get("question", "").strip()
    if not question:
        raise HTTPException(status_code=400, detail="Question is required")

    context = body.get("context", {})
    ui_context = body.get("ui_context", {}) or {}
    history = body.get("history", [])
    screen = ui_context.get("screen") or "upload"

    # Derive provision range dynamically from taxonomy (never hardcode)
    from cam.adapters.lease_review.lease_provision_taxonomy import PROVISIONS
    _std_ids = [p["id"] for p in PROVISIONS if p["id"] != "LP-00" and p.get("default_enabled", True)]
    _prov_range = f"{_std_ids[0]} through {_std_ids[-1]}" if _std_ids else "LP-01 through LP-32"
    _prov_count = len(_std_ids)

    ui_context_lines = []
    if ui_context:
        screen = ui_context.get("screen")
        if screen:
            ui_context_lines.append(f"Current screen: {screen}")
        if ui_context.get("mode"):
            ui_context_lines.append(f"Mode: {ui_context.get('mode')}")
        if ui_context.get("perspective"):
            ui_context_lines.append(f"Perspective: {ui_context.get('perspective')}")
        if "template_loaded" in ui_context:
            ui_context_lines.append(f"Reference file loaded: {'yes' if ui_context.get('template_loaded') else 'no'}")
        if "reference_ready" in ui_context:
            ui_context_lines.append(f"Reference lease ready: {'yes' if ui_context.get('reference_ready') else 'no'}")
        if "tenant_count" in ui_context:
            ui_context_lines.append(f"Tenant file count: {ui_context.get('tenant_count')}")
        if "selected_review_area_count" in ui_context:
            ui_context_lines.append(f"Selected review areas: {ui_context.get('selected_review_area_count')}")
        if "review_areas_ready" in ui_context:
            ui_context_lines.append(f"Review Areas panel ready: {'yes' if ui_context.get('review_areas_ready') else 'no'}")
        if ui_context.get("missing_requirements"):
            ui_context_lines.append("Missing requirements:")
            for item in ui_context.get("missing_requirements") or []:
                ui_context_lines.append(f"- {item}")
        if ui_context.get("next_step"):
            ui_context_lines.append(f"Recommended next step: {ui_context.get('next_step')}")
        if "analyze_enabled" in ui_context:
            ui_context_lines.append(f"Review Leases enabled: {'yes' if ui_context.get('analyze_enabled') else 'no'}")
        if ui_context.get("job_status"):
            ui_context_lines.append(f"Job status: {ui_context.get('job_status')}")
        if "completed_contract_count" in ui_context:
            ui_context_lines.append(f"Completed contracts: {ui_context.get('completed_contract_count')}")
        if "remaining_contract_count" in ui_context:
            ui_context_lines.append(f"Remaining contracts: {ui_context.get('remaining_contract_count')}")
        views = ui_context.get("available_views") or []
        if views:
            ui_context_lines.append("Available views:")
            for view in views:
                label = view.get("label") or view.get("id") or "View"
                purpose = view.get("purpose") or ""
                ui_context_lines.append(f"- {label}: {purpose}".rstrip(": "))
    ui_context_block = "\n".join(ui_context_lines)
    ui_context_prompt = f"UI CONTEXT:\n{ui_context_block}\n\n" if ui_context_block else ""

    screen_intro = (
        "You are a lease analysis guidance assistant for CAM, a commercial lease "
        "deviation and coverage analyzer. The user is on the processing screen while CAM reviews "
        "the uploaded lease or leases. Depending on the selected mode, CAM may be comparing "
        "against a reference lease or running single-document coverage analysis.\n\n"
        "ON THIS SCREEN, PRIORITIZE:\n"
        "- Explaining what CAM is doing now in plain language\n"
        "- Setting expectations for what the user will see when results are ready\n"
        "- Helping the user understand whether they can review partial results yet\n"
        "- Answering broader lease, provision, and review-area questions while they wait\n\n"
        if screen == "processing"
        else
        "You are a lease analysis guidance assistant for CAM, a commercial lease "
        "deviation and coverage analyzer. The user is on the upload page preparing "
        "to review leases in either comparison mode or single-document mode.\n\n"
        "ON THIS SCREEN, PRIORITIZE:\n"
        "- Explaining what CAM does with the uploaded leases\n"
        "- Helping the user understand what to upload and which review areas to keep selected\n"
        "- Telling the user what to do before they click 'Review Leases'\n"
        "- Answering broader lease and review-area questions before analysis starts\n\n"
    )

    system_prompt = (
        screen_intro +
        "THE CURRENT UI FLOW (describe this accurately when asked how to get started):\n"
        "Important: the page may visually default to Compare to reference for layout, but if "
        "UI CONTEXT says the mode was not explicitly selected by the user, do not treat that "
        "default as the user's choice. When asked what to do, briefly explain both paths: "
        "for comparison, upload a reference lease and tenant lease; for single-document "
        "analysis, choose a perspective and upload the lease to analyze.\n\n"
        "Step 1 - Choose a mode:\n"
        "  - Compare to reference: compare one or more tenant leases against a reference/template lease.\n"
        "  - Analyze single document: review one lease without a reference template for coverage gaps and structural issues.\n\n"
        "Step 2 depends on mode:\n"
        "  - Compare to reference: upload the reference lease first.\n"
        "  - Analyze single document: choose the review perspective (Tenant, Landlord, or Neutral / commercially reasonable).\n\n"
        "Step 3 - Upload the lease or leases to review:\n"
        "  - Compare to reference: upload at least one tenant lease after the reference lease is ready.\n"
        "  - Analyze single document: upload at least one lease to analyze.\n\n"
        "Review Areas panel:\n"
        f"  {_prov_range} are checked by default as standard review areas ({_prov_count} total). "
        "The user can uncheck review areas they don't want analyzed, or add custom review areas.\n"
        "  CAM will also automatically surface additional substantive clauses or issue areas it notices during analysis.\n\n"
        "Final step - Click 'Review Leases'.\n\n"

        "WHAT YOU HELP WITH:\n"
        "- Walking users through the steps above when they ask how to get started\n"
        "- Explaining what each review area means and why it matters\n"
        "- Advising which review areas to keep selected for specific lease types\n"
        "- Explaining that CAM automatically discovers additional substantive clauses "
        "or issue areas during analysis and surfaces them in the results\n"
        "- Answering general commercial lease questions\n\n"

        "WHAT YOU CAN AND CANNOT SEE:\n"
        "You CAN see (when provided in context):\n"
        f"- The list of standard review areas ({_prov_range}) the user has selected.\n"
        "- The filenames of uploaded documents.\n"
        "- Any custom review areas the user has added.\n\n"
        "You CANNOT see:\n"
        "- The raw text of the uploaded lease documents.\n"
        "- The full analysis results (those are only available after analysis is complete).\n\n"
        "When asked about the specific text of a clause:\n"
        "- Explain that the actual clause text will be visible in the findings cards once "
        "they run the full analysis — each finding card shows the extracted provision text "
        "from both the template and the tenant lease side by side.\n"
        "- Do NOT say 'I don't have access to your documents' as the only answer — that's "
        "unhelpful. Tell them what you DO know and what they'll see after analysis.\n\n"
        "CRITICAL — UI ACCURACY:\n"
        "- Never reference buttons, links, icons, or interface elements that you are not "
        "certain exist. Do not invent 'View Source', 'See Context', 'Expand', or any "
        "other UI element to direct the user to.\n"
        "- The real UI elements you can reference:\n"
        f"  - {_prov_range} checkboxes in the Standard Review Areas list\n"
        "  - The 'Review Leases' button\n"
        "  - The 'Add custom review area' field at the bottom of the Review Areas panel\n"
        "  - The chat panel (where the user is currently talking to you)\n"
        "- If you are not sure whether a UI element exists, describe the concept "
        "(e.g., 'once you run the analysis, each finding will show the extracted text') "
        "rather than naming a specific button or link.\n\n"

        "WHAT MAKES CAM UNIQUE (use this when asked 'why is CAM unique' or similar):\n"
        "CAM stands for Constrained Assertion Method. Most AI tools produce a single "
        "confident answer even when the evidence is weak or ambiguous — they are "
        "architected to always give you something. CAM is built differently.\n\n"
        "CAM's core approach:\n"
        "1. Multiple independent AI models (Gemini, Claude, GPT, Grok) each evaluate "
        "every provision separately, without seeing each other's answers.\n"
        "2. CAM preserves disagreement rather than hiding it. If models split on whether "
        "a clause is a deviation, that split is surfaced — not collapsed into a majority vote.\n"
        "3. CAM withholds assertions when confidence is insufficient. If the evidence "
        "does not support a clear finding, CAM says so rather than guessing.\n"
        "4. Every finding is backed by a structured reasoning chain — you can see why "
        "a deviation was flagged, not just that it was flagged.\n"
        "5. Severity (Critical, High, Medium, Low) is assessed separately from detection — "
        "a provision can be flagged as a deviation but assessed as low risk in context.\n\n"
        "In practice: CAM catches subtle deviations that single-model tools miss "
        "(like a word change that shifts the legal threshold), flags findings where "
        "models genuinely disagree so a lawyer can make the call, and never gives "
        "false certainty. The goal is to surface what matters, not to fill a report.\n\n"
        "CAM has been validated across scientific fact verification, legal contract "
        "entailment, and commercial lease analysis — the same multi-model disagreement "
        "framework runs underneath all of them.\n\n"

        f"CAM KNOWLEDGE BASE (use this to answer any questions about CAM, scores, signals, pipeline stages, or provisions):\n{CAM_KNOWLEDGE}\n\n"
    )

    if ui_context_prompt:
        system_prompt += ui_context_prompt

    # Add upload context if available
    selected = context.get("selected_review_areas") or context.get("selected_provisions", [])
    if selected:
        system_prompt += f"USER'S SELECTED REVIEW AREAS: {', '.join(selected)}\n"
    uploaded = context.get("uploaded_files", [])
    if uploaded:
        system_prompt += f"UPLOADED FILES: {', '.join(uploaded)}\n"
    custom = context.get("custom_review_areas") or context.get("custom_provisions", [])
    if custom:
        system_prompt += f"CUSTOM REVIEW AREAS ADDED: {', '.join(custom)}\n"

    system_prompt += (
        "\nRESPONSE GUIDELINES:\n"
        "- Be helpful and conversational. This is a guidance tool, not legal advice.\n"
        "- Lead with two things whenever possible: what CAM is doing here, and what the user should do or expect next.\n"
        "- When asked 'how do I get started' or similar, describe the actual flow above concisely.\n"
        "- 100-200 words unless more detail is requested.\n"
        "- Reference provision IDs (LP-XX) when relevant.\n"
        "- No markdown code blocks or code fences. Use plain text formatting.\n"
        "- One brief disclaimer at the end if relevant — never at the start.\n"
        "- If UI CONTEXT is provided, use the exact labels and screen state from it.\n"
        "- Only reference views or interface elements that appear in UI CONTEXT.\n"
        "- When the user asks what to do next, base the answer on UI CONTEXT first.\n"
        "- On the processing screen, be explicit about whether results may still be in progress and what the next visible milestone is.\n"
        "- On the upload screen, be explicit about what is still missing before analysis can start when UI CONTEXT indicates it.\n"
        "- Never start with 'That's a great question' or similar affirmations.\n"
        "- When asked why CAM is unique, explain the multi-model disagreement approach "
        "in plain language. Do NOT describe it as just 'deviation analysis'.\n"
    )

    # Build conversation context from history
    conversation_context = ""
    if history:
        recent = history[-10:]
        for h in recent:
            role = h.get("role", "user")
            content = h.get("content", "")
            if role == "user":
                conversation_context += f"\nUser asked: {content}\n"
            else:
                conversation_context += f"\nAI responded: {content}\n"

    if conversation_context:
        full_prompt = (
            f"CONVERSATION HISTORY:\n{conversation_context}\n"
            f"---\n"
            f"CURRENT QUESTION:\n{question}"
        )
    else:
        full_prompt = question

    try:
        from cam.core.llm import call_llm
        result = _call_chat_with_followups(
            call_llm=call_llm,
            provider="google",
            system_prompt=system_prompt,
            user_prompt=full_prompt,
            temperature=0.3,
        )
        return {
            "response": result.get("content", ""),
            "suggested_followups": result.get("suggested_followups", []),
            "model": result.get("model", ""),
        }
    except Exception as e:
        logger.error(f"General chat error: {e}")
        raise HTTPException(status_code=500, detail=f"Chat error: {e}")


@app.delete("/api/jobs/{job_id}")
async def delete_job(job_id: str):
    """Manually delete a job and all its results/uploads."""
    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    job_manager.delete_job(job_id)
    return {"status": "deleted", "job_id": job_id}


@app.get("/api/models")
def get_models():
    """Return model display names for the frontend — single source of truth."""
    from cam.adapters.lease_review.model_config import DISPLAY_NAMES, CHAT_DEFAULTS
    return {
        "display_names": DISPLAY_NAMES,
        "chat_defaults": {
            k: v[1] for k, v in CHAT_DEFAULTS.items()
        },
    }


@app.get("/api/provisions")
def get_provisions():
    """Return all default provisions for the frontend checklist."""
    from cam.adapters.lease_review.lease_provision_taxonomy import PROVISIONS
    return [
        {
            "id": p["id"],
            "name": p["name"],
            "description": p["description"],
            "default_enabled": p.get("default_enabled", True),
            "always_on": p.get("always_on", False),
            "identity_check": p.get("identity_check", False),
        }
        for p in PROVISIONS
    ]


# ── Summary document generation endpoints ──

@app.get("/api/jobs/{job_id}/results/{tenant_index}/summary")
def download_tenant_summary(job_id: str, tenant_index: int):
    """Generate and download a per-tenant summary DOCX document."""
    from app.summary_generator import generate_tenant_summary

    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job.get("status") == "completed" and job_manager.is_job_expired(job):
        return JSONResponse(status_code=410, content={"detail": "This analysis has expired."})

    if job["status"] != "completed":
        raise HTTPException(status_code=404, detail="Job not yet complete")

    tenants = job.get("input_config", {}).get("tenants", [])
    if tenant_index < 0 or tenant_index >= len(tenants):
        raise HTTPException(status_code=404, detail="Invalid tenant index")

    tenant = tenants[tenant_index]
    result_path = tenant.get("result_path")
    if not result_path or not Path(result_path).exists():
        raise HTTPException(status_code=404, detail="Results not available for this tenant")

    try:
        pipeline_results = json.loads(Path(result_path).read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as e:
        raise HTTPException(status_code=500, detail=f"Failed to read results: {e}")

    config = get_config()
    output_dir = Path(config["RESULTS_DIR"]) / job_id
    output_dir.mkdir(parents=True, exist_ok=True)
    tenant_name = Path(tenant["filename"]).stem
    output_path = str(output_dir / f"Summary_{tenant_name}.docx")

    try:
        generate_tenant_summary(
            pipeline_results,
            output_path,
            resolutions=job.get("resolutions", {}),
            tenant_idx=tenant_index,
        )
    except Exception as e:
        logger.error(f"Summary generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Summary generation failed: {e}")

    return FileResponse(
        path=output_path,
        filename=f"Summary_{tenant_name}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/api/jobs/{job_id}/summary")
def download_batch_summary(job_id: str):
    """Generate and download a combined summary PDF covering all tenants."""
    from app.summary_generator import generate_combined_summary_pdf

    job = job_manager.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if job.get("status") == "completed" and job_manager.is_job_expired(job):
        return JSONResponse(status_code=410, content={"detail": "This analysis has expired."})

    if job["status"] != "completed":
        raise HTTPException(status_code=404, detail="Job not yet complete")

    tenants = job.get("input_config", {}).get("tenants", [])
    tenant_results = []
    for tenant in tenants:
        result_path = tenant.get("result_path")
        if result_path and Path(result_path).exists():
            try:
                data = json.loads(Path(result_path).read_text(encoding="utf-8"))
                tenant_results.append(data)
            except (json.JSONDecodeError, OSError):
                pass

    if not tenant_results:
        raise HTTPException(status_code=404, detail="No tenant results available")

    config = get_config()
    output_dir = Path(config["RESULTS_DIR"]) / job_id
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = str(output_dir / "Lease_Analysis_Synopsis.pdf")

    try:
        generate_combined_summary_pdf(job, tenant_results, output_path)
    except Exception as e:
        logger.error(f"Summary PDF generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Summary generation failed: {e}")

    return FileResponse(
        path=output_path,
        filename="Lease_Analysis_Synopsis.pdf",
        media_type="application/pdf",
    )


# ── Results page route (frontend will handle rendering) ──

@app.get("/results/{job_id}")
def serve_results_page(job_id: str):
    """Serve the frontend for viewing results (same index.html, JS handles routing)."""
    index_path = static_dir / "index.html"
    if index_path.exists():
        return FileResponse(
            str(index_path),
            headers={"Cache-Control": "no-cache, no-store, must-revalidate"},
        )
    return JSONResponse({"message": "Results page — frontend not yet built", "job_id": job_id})
