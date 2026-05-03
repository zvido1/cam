"""
CAM Lease Review — DOCX Annotator

Inserts colored callout text blocks on deviating/unclear provisions in a
tenant DOCX file.  Callout blocks are inserted immediately after the
paragraph containing the flagged provision text.

Note: Native Word comment insertion via XML manipulation was attempted but
proved incompatible with python-docx 0.8.11's Part/OPC structure (corrupts
the document on save).  Callout blocks are the reliable approach.
"""

import re
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _format_comment_text(provision: dict, resolution: dict = None) -> str:
    """Format the comment text for a provision finding."""
    pid = provision.get("provision_id", "?")
    pname = provision.get("provision_name", "")
    verdict = provision.get("final_verdict", "?")
    severity = provision.get("severity", "")
    pattern = provision.get("agreement_pattern", "")

    lines = [f"[CAM \u2014 {pid} {pname}]"]
    lines.append(f"Status: {verdict} ({severity})")
    lines.append(f"Agreement: {pattern}")
    lines.append("")

    # Challenge details or cascade mechanism
    cascade_src = provision.get("cascade_source")
    if provision.get("cascade_verdict") == "CASCADE_MATERIAL":
        if cascade_src and cascade_src.get("term"):
            lines.append(f"Definition cascade from: \"{cascade_src['term']}\"")
        lines.append(f"Definition cascade: {provision.get('cascade_mechanism', '')}")
        lines.append(f"Impact: {provision.get('cascade_impact', '')}")
    elif provision.get("challenge_details"):
        lines.append(provision["challenge_details"])
    elif provision.get("severity_reasoning"):
        lines.append(provision["severity_reasoning"])

    lines.append("")

    # Fragility signals
    frag = provision.get("fragility", {})
    if frag.get("signals"):
        lines.append(f"Fragility: {', '.join(frag['signals'])}")

    # Recommended action
    action = provision.get("recommended_action", "")
    if action and action != "no_action":
        action_labels = {
            "note_for_awareness": "Note for awareness",
            "attorney_review_recommended": "Attorney review recommended",
            "attorney_review_required": "Attorney review required",
        }
        lines.append(f"\u2192 {action_labels.get(action, action)}")

    # Lawyer's resolutions from CAM review session
    if resolution:
        lines.append("")
        lines.append("\u2014 Lawyer's Review \u2014")
        status = resolution.get("status", "")
        status_labels = {
            "accepted": "Accepted as-is",
            "needs_negotiation": "Needs Negotiation",
            "not_applicable": "Not Applicable",
            "resolved": "Resolved",
        }
        if status:
            lines.append(f"Decision: {status_labels.get(status, status)}")
        concern = resolution.get("concern_state", "")
        concern_reason = resolution.get("concern_reason", "")
        if concern == "flagged" and concern_reason:
            lines.append(f"Concern: {concern_reason}")
        notes = resolution.get("notes", [])
        if notes:
            for note in notes:
                text = note.get("text", "") if isinstance(note, dict) else str(note)
                if text:
                    lines.append(f"Note: {text}")

    return "\n".join(lines)


def _add_callout_block(document, paragraph, comment_text, severity="MEDIUM"):
    """Fallback: Insert a colored text block after the paragraph."""
    # Determine color hex based on severity
    color_hex_map = {
        "CRITICAL": "CC0000",   # Red
        "HIGH": "CC4400",       # Dark orange
        "MEDIUM": "CC8800",     # Orange
        "LOW": "888800",        # Dark yellow
        "REVIEW": "0066CC",     # Blue
        "CONFORMS": "CC8800",   # Default orange
    }
    color_hex = color_hex_map.get(severity, "CC8800")

    # Find the paragraph's parent and position
    parent = paragraph._element.getparent()
    idx = list(parent).index(paragraph._element)

    # Create a new paragraph for the callout
    callout_para = OxmlElement("w:p")

    # Add border/shading to make it look like a callout
    ppr = OxmlElement("w:pPr")
    # Add shading (light background)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF3CD")  # Light yellow background
    ppr.append(shd)
    # Add borders
    borders = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "1")
        border.set(qn("w:color"), color_hex)
        borders.append(border)
    ppr.append(borders)
    callout_para.append(ppr)

    # Add the text
    for line in comment_text.split("\n"):
        if line:
            run = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "18")  # 9pt font
            rpr.append(sz)
            clr = OxmlElement("w:color")
            clr.set(qn("w:val"), color_hex)
            rpr.append(clr)
            run.append(rpr)
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = line
            run.append(t)
            callout_para.append(run)

            # Add line break
            br_run = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br_run.append(br)
            callout_para.append(br_run)

    # Insert after the target paragraph
    parent.insert(idx + 1, callout_para)


def _find_paragraph_by_text(document, search_text: str, min_match_chars: int = 40):
    """Find a paragraph containing the search text (first ~N chars).

    Returns the first matching paragraph or None.
    """
    if not search_text:
        return None

    # Use first N chars, cleaned up
    needle = search_text[:min_match_chars].strip()
    if len(needle) < 10:
        needle = search_text.strip()
    if not needle:
        return None

    needle_lower = needle.lower()

    for para in document.paragraphs:
        if needle_lower in para.text.lower():
            return para

    # Fallback: try with fewer characters
    if len(needle) > 20:
        short_needle = needle[:20].lower()
        for para in document.paragraphs:
            if short_needle in para.text.lower():
                return para

    return None


def _insert_summary_section(doc, results):
    """Insert a summary section at the top of the DOCX before annotations."""
    from datetime import datetime
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    body = doc.element.body
    insert_before = body[0] if len(body) > 0 else None

    def _add_para(text, size=9, bold=False, color=None, space_after=2,
                   shading=None, border_color=None, indent=0):
        """Add a paragraph. shading=fill hex, border_color=left border hex."""
        p = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), str(space_after * 20))
        ppr.append(spacing)
        if indent:
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(indent))
            ppr.append(ind)
        if shading:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), shading)
            ppr.append(shd)
        if border_color:
            pbdr = OxmlElement("w:pBdr")
            for side in ["top", "left", "bottom", "right"]:
                bd = OxmlElement(f"w:{side}")
                bd.set(qn("w:val"), "single")
                bd.set(qn("w:sz"), "4" if side != "left" else "12")
                bd.set(qn("w:space"), "1")
                bd.set(qn("w:color"), border_color if side == "left" else "E5E7EB")
                pbdr.append(bd)
            ppr.append(pbdr)
        p.append(ppr)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size * 2))
        rpr.append(sz)
        if bold:
            b = OxmlElement("w:b")
            rpr.append(b)
        if color:
            c = OxmlElement("w:color")
            c.set(qn("w:val"), color)
            rpr.append(c)
        run.append(rpr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        run.append(t)
        p.append(run)
        if insert_before is not None:
            body.insert(list(body).index(insert_before), p)
        else:
            body.append(p)

    # Header
    _add_para("CAM Lease Analysis Report", size=16, bold=True, color="1A365D", space_after=4)
    date_str = datetime.now().strftime("%B %d, %Y")
    tenant_file = results.get("tenant_file", "")
    _add_para(f"{date_str}  |  {tenant_file}", size=9, color="64748B", space_after=8)

    # Contract metadata
    meta = results.get("contract_metadata", {})
    has_meta = False
    for label, key in [("Landlord", "landlord"), ("Property", "property_description"),
                       ("Term", "term_length"), ("Base Rent", "base_rent"),
                       ("Tenant", "tenant"), ("Governing Law", "governing_law")]:
        val = meta.get(key, "")
        if val and len(str(val).replace("_", "").replace("TBD", "").strip()) >= 5:
            if not has_meta:
                _add_para("Contract Summary", size=12, bold=True, color="1A365D", space_after=2)
                has_meta = True
            _add_para(f"{label}: {val}", size=9)
    if has_meta:
        _add_para("", size=4, space_after=6)

    # Provisions checklist (traffic light grouping by severity)
    provisions = results.get("provisions", [])
    if provisions:
        _add_para("Provisions Analyzed", size=12, bold=True, color="1A365D", space_after=2)

        sev_groups = {"CRITICAL": [], "HIGH": [], "MEDIUM": [], "LOW": [], "CONFORMS": []}
        for p in provisions:
            if p.get("final_verdict") == "DEVIATES":
                sev = p.get("severity", "MEDIUM")
                sev_groups.get(sev, sev_groups["MEDIUM"]).append(p)
            else:
                sev_groups["CONFORMS"].append(p)

        sev_colors = {
            "CRITICAL": "DC2626", "HIGH": "C2410C", "MEDIUM": "D97706",
            "LOW": "64748B", "CONFORMS": "16A34A",
        }
        sev_markers = {
            "CRITICAL": "\u2715 CRITICAL",   # ✕
            "HIGH":     "\u25CF HIGH",        # ●
            "MEDIUM":   "\u25CB MEDIUM",      # ○
            "LOW":      "\u2013 LOW",         # –
            "CONFORMS": "\u2713 CONFORMS",   # ✓
        }
        sev_shading = {
            "CRITICAL": "FEF2F2", "HIGH": "FFF7ED", "MEDIUM": "FFFBEB",
            "LOW": "F8FAFC", "CONFORMS": "F0FDF4",
        }
        for sev_key in ["CRITICAL", "HIGH", "MEDIUM", "LOW", "CONFORMS"]:
            items = sev_groups[sev_key]
            if not items:
                continue
            names = ", ".join(
                f"{p.get('provision_id', '')} {(p.get('provision_name', '') or '').split(' ', 1)[-1]}"
                for p in items
            )
            _add_para(
                f"  {sev_markers[sev_key]}: {names}",
                size=9, color=sev_colors[sev_key],
                shading=sev_shading[sev_key],
                border_color=sev_colors[sev_key],
                space_after=1,
            )

        _add_para("", size=4, space_after=6)

    # Findings
    deviations = [p for p in provisions if p.get("final_verdict") == "DEVIATES"]
    sev_order = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}
    deviations.sort(key=lambda d: sev_order.get(d.get("severity", "LOW"), 99))

    if deviations:
        _add_para("Findings", size=12, bold=True, color="1A365D", space_after=4)
        action_labels = {
            "note_for_awareness": "Note for awareness",
            "attorney_review_recommended": "Attorney review recommended",
            "attorney_review_required": "Attorney review required",
        }
        for d in deviations:
            pid = d.get("provision_id", "")
            pname = d.get("provision_name", "")
            sev = d.get("severity", "")
            headline = d.get("risk_headline", "")
            what_changed = d.get("challenge_details", d.get("what_changed", ""))
            impact = d.get("severity_reasoning", "")
            financial = d.get("financial_impact", "")
            action = d.get("recommended_action", "")
            cascade_src = d.get("cascade_source")
            action_text = action_labels.get(action, action)
            verdicts = d.get("evaluator_verdicts", {})
            agreeing = sum(1 for v in verdicts.values() if v == "DEVIATES")
            total = len(verdicts) or 3

            sev_clr = "DC2626" if sev == "CRITICAL" else "C2410C" if sev == "HIGH" else "D97706" if sev == "MEDIUM" else "64748B"
            _add_para(f"{pid} {pname}  [{sev}]  ({agreeing}/{total} evaluators agree)", size=10, bold=True, color=sev_clr)

            # Cascade source (if present)
            if cascade_src and cascade_src.get("term"):
                cs_text = f"Definition cascade from: \"{cascade_src['term']}\""
                if cascade_src.get("defined_in"):
                    cs_text += f" -- {cascade_src['defined_in']}"
                _add_para(cs_text, size=8, color="D97706")

            if headline:
                _add_para(headline, size=9, bold=True, color=sev_clr)
            if what_changed:
                _add_para(f"What Changed: {what_changed}", size=9)
            if impact:
                _add_para(f"Impact: {impact}", size=9)
            if financial:
                _add_para(f"Financial Impact: {financial}", size=9)
            if action_text and action_text != "no_action":
                _add_para(f"Recommended Action: {action_text}", size=9, color="1A365D")
            _add_para("", size=4, space_after=6)

    # Step 297c: Provision Conflicts section
    conflicts = results.get("conflicts", []) or []
    if conflicts:
        _add_para("Provision Conflicts", size=12, bold=True, color="92400E", space_after=4)
        _add_para(
            "The following pairs of provisions create internal conflicts within the lease.",
            size=9, color="78350F", space_after=4,
        )
        conflict_sev_colors = {"high": "B91C1C", "medium": "D97706", "low": "6B7280"}
        conflict_shading   = {"high": "FEF2F2", "medium": "FFF7ED", "low": "F8FAFC"}
        for c in conflicts:
            cid = c.get("id", "")
            cname = c.get("name", "")
            sev = c.get("severity", "medium")
            lps = ", ".join(c.get("lps_implicated", []) or [])
            desc = c.get("description", "")
            sev_label = sev.upper()
            clr = conflict_sev_colors.get(sev, "6B7280")
            shd = conflict_shading.get(sev, "F8FAFC")
            _add_para(
                f"[CONFLICT — {sev_label}] {cid}: {cname}",
                size=10, bold=True, color=clr, shading=shd, border_color=clr, space_after=1,
            )
            if lps:
                _add_para(f"Implicates {lps}.", size=9, indent=180)
            if desc:
                _add_para(desc, size=9, indent=180)
            _add_para("", size=4, space_after=4)

    # Disclaimer + divider. Step 277: Mode A uses the multi-evaluator
    # adjudication pipeline; Mode C uses a single-model coverage layer.
    # Pick the phrasing that matches what actually generated this run.
    is_mode_c = (results.get("mode") == "analyze")
    if is_mode_c:
        disclaimer_body = (
            "This analysis was generated by CAM's coverage analysis layer. "
            "It is intended as a review aid, not legal advice. All findings should "
            "be verified by qualified legal counsel before any action is taken."
        )
    else:
        disclaimer_body = (
            "This analysis was generated by the CAM (Constrained Assertion Method) system using multiple "
            "independent AI evaluators. It is intended as a review aid, not legal advice. All findings should "
            "be verified by qualified legal counsel before any action is taken."
        )
    _add_para("---", size=8, color="94A3B8", space_after=4)
    _add_para(disclaimer_body, size=7, color="94A3B8", space_after=10)
    _add_para("--- End of CAM Summary ---", size=8, bold=True, color="94A3B8", space_after=12)


def _format_coverage_callout_text(coverage_item: dict, cov_resolution: dict = None,
                                  perspective: str = "tenant") -> str:
    """Format the callout text for a Mode C coverage gap finding.

    Mirrors `lease_pdf_annotator._format_coverage_annotation_text` so DOCX
    and PDF coverage callouts read the same. Step 273: state_label is now
    perspective-aware via `_resolve_display`.
    """
    from cam.adapters.lease_review.lease_display import extract_headline

    pid = coverage_item.get("issue_area_id", "?")
    # Step 279: prefer the schema's clean issue_area_name over
    # provision_name (which is None in Mode C runs).
    pname = (coverage_item.get("provision_name")
             or coverage_item.get("issue_area_name")
             or "")
    materiality = coverage_item.get("materiality", "medium")
    exposure = coverage_item.get("exposure_statement", "")
    elements_missing = coverage_item.get("elements_missing", [])
    # Step 278: integrate headline into the header line.
    headline = (coverage_item.get("exposure_headline") or "").strip()
    if not headline and exposure:
        headline = extract_headline(exposure)

    mat_labels = {"high": "HIGH", "medium": "MEDIUM", "low": "LOW"}
    mat_label = mat_labels.get(materiality, materiality.upper())

    # Step 279: single-line header. Marker [GAP] carries the gap-bucket
    # signal; materiality stays as a parenthetical at the end (high vs
    # medium vs low isn't carried by the marker). The state label
    # (UNFAVORABLE TERMS / FAVORABLE TERMS / etc.) was redundant with
    # the marker + materiality and is dropped.
    if headline:
        lines = [f"[GAP] {pid} {pname} — {headline} ({mat_label} materiality)"]
    else:
        lines = [f"[GAP] {pid} {pname} ({mat_label} materiality)"]
    lines.append("")

    if elements_missing:
        missing_str = ", ".join(str(e) for e in elements_missing[:5])
        lines.append(f"Missing: {missing_str}")
        lines.append("")

    if exposure:
        lines.append(exposure)

    if cov_resolution:
        status = cov_resolution.get("status", "")
        status_labels = {
            "reviewed": "Reviewed",
            "flagged": "Flagged for follow-up",
            "accepted": "Risk accepted",
        }
        if status and status != "open":
            lines.append("")
            lines.append("— Lawyer's Review —")
            lines.append(f"Decision: {status_labels.get(status, status)}")
        notes = cov_resolution.get("notes", []) or []
        for note in notes:
            text = note.get("text", "") if isinstance(note, dict) else str(note)
            if text:
                lines.append(f"Note: {text}")

    result = "\n".join(line for line in lines if line is not None)
    # Strip markdown bold/italic markers (mirrors deviation annotation handling)
    result = re.sub(r'\*\*(.+?)\*\*', r'\1', result)
    result = re.sub(r'\*(.+?)\*', r'\1', result)
    return result


# Distinct purple palette for coverage gap callouts. Mirrors the PDF
# annotator's `_coverage_color` choice so the two formats are visually
# consistent. The hex border color is `7B5EAB`, with a light-lavender fill.
_COVERAGE_BORDER_HEX = "7B5EAB"
_COVERAGE_FILL_HEX = "F1ECFA"


def _add_coverage_callout_block(document, paragraph, comment_text):
    """Insert a purple coverage-gap callout block after the paragraph.

    Mirrors `_add_callout_block` but uses the coverage palette so the gap
    callouts read as distinct from deviation callouts in Mode A documents
    where both can appear on the same provision.
    """
    parent = paragraph._element.getparent()
    idx = list(parent).index(paragraph._element)

    callout_para = OxmlElement("w:p")

    ppr = OxmlElement("w:pPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _COVERAGE_FILL_HEX)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "1")
        border.set(qn("w:color"), _COVERAGE_BORDER_HEX)
        borders.append(border)
    ppr.append(borders)
    callout_para.append(ppr)

    for line in comment_text.split("\n"):
        if line:
            run = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "18")  # 9pt
            rpr.append(sz)
            clr = OxmlElement("w:color")
            clr.set(qn("w:val"), _COVERAGE_BORDER_HEX)
            rpr.append(clr)
            run.append(rpr)
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = line
            run.append(t)
            callout_para.append(run)

            br_run = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br_run.append(br)
            callout_para.append(br_run)

    parent.insert(idx + 1, callout_para)


def annotate_docx(
    original_docx_path: str,
    results: dict,
    output_path: str,
    resolutions: dict = None,
    cov_resolutions: dict = None,
) -> str:
    """Insert Word comments on deviating provisions in tenant DOCX.

    Args:
        original_docx_path: Path to the original tenant DOCX file.
        results: Pipeline results dict (with "provisions" list).
        output_path: Path for the annotated output DOCX.

    Returns:
        Path to the annotated DOCX file.
    """
    doc = Document(original_docx_path)

    # Prepend summary section at top of document
    try:
        _insert_summary_section(doc, results)
    except Exception as e:
        print(f"[docx_annotator] Summary section insertion failed (non-fatal): {e}", flush=True)

    annotations_added = 0
    not_found = 0

    for provision in results.get("provisions", []):
        if provision.get("final_verdict") not in ("DEVIATES", "UNCLEAR"):
            continue

        # Look up resolution for this provision (keyed by tenant_idx:provision_id)
        resolution = None
        if resolutions:
            pid = provision.get("provision_id", "")
            # Try both keying formats used by job_manager
            resolution = (resolutions.get(f"0:{pid}") or
                          resolutions.get(pid) or
                          None)
        comment_text = _format_comment_text(provision, resolution)
        search_text = provision.get("tenant_text", "")

        # Find the paragraph containing this provision
        para = _find_paragraph_by_text(doc, search_text)

        if para is None:
            # Try searching by provision name in section refs
            section_ref = provision.get("tenant_section_ref", "")
            if section_ref:
                para = _find_paragraph_by_text(doc, section_ref, min_match_chars=15)

        if para is None:
            pid = provision.get("provision_id", "?")
            print(f"[docx_annotator] Could not locate text for {pid}", flush=True)
            not_found += 1
            continue

        _add_callout_block(doc, para, comment_text, provision.get("severity", "MEDIUM"))
        annotations_added += 1

    # === Coverage gap callouts (Step 271, perspective-aware Step 273) ===
    # Drop [GAP] callout blocks for provisions whose display bucket is one
    # of the surfaced annotation buckets (needs_attention,
    # favorable_to_your_side, asymmetric_terms, worth_reviewing). Items
    # entirely missing from the document have no anchor and are skipped
    # (same as the PDF annotator) — they still appear in the Synopsis.
    from cam.adapters.lease_review.lease_display import (
        _resolve_display, ANNOTATED_BUCKETS, resolve_perspective,
    )

    perspective = resolve_perspective(results)
    coverage_assessment = results.get("coverage_assessment", []) or []
    provisions_by_id = {p.get("provision_id"): p for p in results.get("provisions", []) or []}

    coverage_callouts_added = 0
    coverage_not_found = 0

    for cov in coverage_assessment:
        disp = _resolve_display(cov, perspective)
        if disp["bucket"] not in ANNOTATED_BUCKETS:
            continue
        state = cov.get("coverage_state", "covered")
        if state == "missing":
            continue

        pid = cov.get("issue_area_id", "")
        if not pid:
            continue

        prov = provisions_by_id.get(pid, {})
        anchor_text = (prov.get("tenant_text", "") or "").strip()
        section_ref = (prov.get("tenant_section_ref", "") or "").strip()
        issue_area_name = (cov.get("issue_area_name") or cov.get("provision_name") or "").strip()

        cov_resolution = None
        if cov_resolutions:
            cov_resolution = (
                cov_resolutions.get(f"cov:0:{pid}")
                or cov_resolutions.get(f"cov:{pid}")
            )

        comment_text = _format_coverage_callout_text(cov, cov_resolution, perspective)

        # Anchor fallback chain mirrors the PDF annotator coverage path:
        # tenant_text -> section_ref -> issue_area_name -> issue_area_id.
        # Mode A typically resolves on tenant_text; Mode C falls through to
        # name/id since provisions[] is empty.
        para = None
        if anchor_text:
            para = _find_paragraph_by_text(doc, anchor_text)
        if para is None and section_ref:
            para = _find_paragraph_by_text(doc, section_ref, min_match_chars=15)
        if para is None and issue_area_name:
            para = _find_paragraph_by_text(doc, issue_area_name, min_match_chars=10)
        if para is None and pid:
            para = _find_paragraph_by_text(doc, pid, min_match_chars=4)

        if para is None:
            coverage_not_found += 1
            print(f"[docx_annotator] Could not anchor coverage gap for {pid}", flush=True)
            continue

        _add_coverage_callout_block(doc, para, comment_text)
        coverage_callouts_added += 1

    # Save the annotated document
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    print(
        f"[docx_annotator] Saved {output_path} "
        f"({annotations_added} deviations, {coverage_callouts_added} coverage gaps, "
        f"{not_found + coverage_not_found} not found)",
        flush=True,
    )

    return output_path
